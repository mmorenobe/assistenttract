import streamlit as st
from itertools import chain
from docx.shared import Pt
from docx.shared import Inches
import os


class AssistentDolencies:
    def __init__(self):
        self.preguntes_inicials = {
            "Nom del pacient": "",
            "És dona o home?": ["Dona", "Home"],
            "Quins símptomes músculoesquelètics té?": "",
            "Explicació dels símptomes": "",
            "Des de quan?": "",
            "En quins moments del dia li passa? Quan o on es troba quan li passa en general?": "",
            "Té lesions vertebrals diagnosticades?": ["Sí", "No"],
            "Té malalties cròniques?": ["Sí", "No"],
            "Té al·lèrgies?": ["Sí", "No"]
        }

   # Diccionari de sinònims
        self.sinònims = {
            "mal de cap": "cefalea",
            "mal de panxa": "dolor abdominal",
            "mal de ronyons": "dolor lumbar",
            "punxades al cap": "cefalea tensional",
            "dolor cervical": "cervicàlgia",
            "tortícoli": "cervicàlgia aguda",
            "espatlla bloquejada": "tendinopatia de l'espatlla",
            "rigidesa de l’esquena": "rigidesa dorsal",
            "punxades al pit": "dolor toràcic",
            "mal d’ossos": "dolor articular",
            "dolor d’estómac": "gastràlgia",
            "cames adolorides": "mialgia a les extremitats inferiors",
            "mal a la pelvis": "dolor pèlvic",
            "panxa inflada": "distensió abdominal",
            "cremor d’estómac": "pirosi",
            "gasos": "meteorisme",
            "reflux": "reflux gastroesofàgic",
            "restrenyiment": "estrenyiment",
            "diarrea": "deposicions líquides",
            "ganes de vomitar": "nàusees",
            "vòmits": "emesi",
            "sentir-se ple": "sensació de plenitud gàstrica",
            "digestions lentes": "dispèpsia",
            "amargor a la boca": "reflux biliar",
            "cansament": "fatiga",
            "molt cansament al matí": "astenia matutina",
            "son després de dinar": "somnolència postprandial",
            "sensació de debilitat": "astenia",
            "falta d'energia": "fatiga crònica",
            "palpitacions": "taquicàrdia",
            "batecs forts del cor": "palpitacions",
            "baixa tensió": "hipotensió",
            "pujades de pressió": "hipertensió",
            "sensació d’ofec": "dispnea",
            "ofec": "dispnea",
            "tos seca": "tos irritativa",
            "mocs al pit": "congestió bronquial",
            "sensació de nas tapat": "congestió nasal",
            "mocs grocs": "rinitis infecciosa",
            "pell greixosa": "seborrea",
            "pell seca i esquerdada": "xerosi",
            "granets": "acné",
            "picor a la pell": "prurit",
            "taques a la pell": "hiperpigmentació",
            "pell groguenca": "icterícia",
            "marejos": "vertigen",
            "cap carregat": "cefalea tensional",
            "falta d’equilibri": "inestabilitat postural",
            "cames formiguejant": "parestèsies a les extremitats inferiors",
            "mans adormides": "parestèsies a les extremitats superiors",
            "mal a les cames": "mialgia a les cames",
            "mal als braços": "mialgia als braços",
            "tibantor als bessons": "contractura muscular als bessons",
            "punxades a la cama": "ciàtica",
            "ganes de fer pipi sovint": "poliúria",
            "fer pipi moltes vegades a la nit": "nictúria",
            "pipí fosc": "orina concentrada",
            "orina amb olor forta": "orina carregada",
            "coïssor en fer pipi": "disúria",
            "regles irregulars": "cicles menstruals irregulars",
            "sang molt fosca a la regla": "alteració de la coloració del sagnat menstrual",
            "regla molt abundant": "hipermenorrea",
            "regla molt escassa": "hipomenorrea",
            "dolor menstrual": "dismenorrea",
            "estrès": "tensió emocional",
            "estar nerviós": "ansietat",
            "sensació d’angoixa": "ansietat generalitzada",
            "preocupació constant": "ruminació mental",
            "no puc dormir": "insomni",
            "molts pensaments al cap": "hiperactivitat mental",
            "tristesa sense motiu": "distímia",
            "ulls grocs": "icterícia ocular",
            "ulls vermells": "hiperèmia conjuntival",
            "ulls secs": "xeroftàlmia",
            "boca seca": "xerostomia",
            "ulls inflamats": "blefaritis",
            "sensació de pressió al cap": "cefalea tensional"
        }
        self.sinònims.update({
            "dolor d'estómac": "gastràlgia",
            "mal de panxa": "dolor abdominal",
            "malestar digestiu": "dispepsia",
            "cervicals carregades": "cervicàlgia",
            "rigidesa cervical": "cervicàlgia",
            "mareig": "vertigen",
            "sensació d'inestabilitat": "vertigen",
            "contractura a l'esquena": "rigidesa dorsal",
            "molèstia lumbar": "dolor lumbar",
            "punxades al genoll": "dolor articular",
            "dolor al genoll": "condromalàcia rotuliana",
            "dolor al gluti": "síndrome del piramidal",
            "cames pesades": "mialgia a les extremitats inferiors",
            "mandra i fatiga": "astenia",
            "boca seca": "xerostomia",
            "ulls irritats": "hiperèmia conjuntival",
            "inflor abdominal": "distensió abdominal",
            "digestions pesades": "dispèpsia",
            "esquena carregada": "rigidesa dorsal"
        })
        # Diccionari de tractaments de Fitoteràpia classificats per categories
        self.classificacio_fitoterapia = {
            # Categories generals
            "Dolor de canell - ganglió a mà dreta": "Dolor canell - ganglió",
            "Dolor de canell - ganglió a mà dreta més faringitis freqüents": "Dolor canell - ganglió",
            "Dolor de canell - ganglió a mà esquerra": "Dolor canell - ganglió",
            "Dolor de canell - ganglió a mà esquerra més faringitis freqüents": "Dolor canell - ganglió",
            "Dolor de canell - ganglió bilateral": "Dolor canell - ganglió",
            "Dolor de canell - ganglió bilateral més faringitis freqüent": "Dolor canell - ganglió",
            "Dolor de canell – ganglió, tractament senzill sense infusions": "Dolor canell - ganglió",
            "Tendinopatia aquil·liana dreta crònica en home": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana dreta crònica en dona": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana dreta amb menstruacions més curtes i/o escasses": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana dreta en home aguda o recent": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana dreta en dona aguda o recent": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana esquerra aguda o recent": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana esquerra crònica en home": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana bilateral en pacient home": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana bilateral en home amb infeccions d'orina prèvies": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana bilateral en pacient dona": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana bilateral en dona amb infeccions d'orina prèvies": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana amb infeccions d'orina recurrents i problema digestiu recurrent (indistintament dreta o esquerra)": "Tendinopatia aquil·liana",
            "Tendinopatia aquil·liana per a algú que no prendrà infusions": "Tendinopatia aquil·liana",
            "Tensió de bessons": "Tensió de bessons",
            "Tensió de bessons més faringitis prèvies sovint": "Tensió de bessons",
            "Tensió de bessons més infeccions d'orina prèvies": "Tensió de bessons",
            "Tensió de bessons més faringitis i infeccions d'orina prèvies": "Tensió de bessons",
            "Tensió de bessons més problemes digestius": "Tensió de bessons",
            "Tensió de bessons més problemes digestius i faringitis freqüents": "Tensió de bessons",
            "Tensió de bessons més problemes digestius i infeccions d'orina": "Tensió de bessons",
            "Tensió de bessons més faringitis i infeccions d'orina prèvies i problemes digestius": "Tensió de bessons",
            "Tensió de bessons tractament senzill sense infusions": "Tensió de bessons",
            "Trencament bessons tractament de base": "Trencament bessons",
            "Trencament de bessons més faringitis prèvies sovint": "Trencament bessons",
            "Trencament de bessons més infeccions d'orina prèvies": "Trencament bessons",
            "Trencament de bessons més faringitis i infeccions d'orina prèvies": "Trencament bessons",
            "Trencament de bessons més problemes digestius": "Trencament bessons",
            "Trencament de bessons més problemes digestius i faringitis freqüents": "Trencament bessons",
            "Trencament de bessons més problemes digestius i infeccions d'orina": "Trencament bessons",
            "Trencament de bessons més faringitis i infeccions d'orina prèvies i problemes digestius": "Trencament bessons",
            "Trencament de bessons, tractament senzill sense infusions": "Trencament bessons",
            "Cervicàlgia esquerra": "Cervicàlgia",
            "Cervicàlgia esquerra amb marejos": "Cervicàlgia",
            "Cervicàlgia esquerra amb cremor d'estómac i gastritis": "Cervicàlgia",
            "Torticolis esquerra": "Cervicàlgia",
            "Cervicàlgia general": "Cervicàlgia",
            "Cervicàlgia general amb marejos": "Cervicàlgia",
            "Pacient amb dorsàlgia esquerra més cremor d'estómac o gastritis": "Dorsàlgia",
            "Pacient amb dorsàlgia esquerra més cremor d'estómac o gastritis ( sense infusions)": "Dorsàlgia",
            "Pacient amb dorsàlgia esquerra sense cremor d'estómac o gastritis ( sense infusions)": "Dorsàlgia",
            "Dorsàlgia general difusa bilateral ( tractament senzill , sense infusions)": "Dorsàlgia",
            "Dorsàlgia central amb irradiació a la boca de l'estómac": "Dorsàlgia",
            "Dorsàlgia general més cervicàlgia bilateral ( tractament senzill , sense infusions)": "Dorsàlgia",
            "Dorsàlgia general més cervicàlgia bilateral": "Dorsàlgia",
            "Síndrome de Tietze esquerre amb cremor d'estómac o gastritis": "Síndrome de Tietze",
            "Síndrome de Tietze esquerre amb cremor d'estómac o gastritis ( sense infusions)": "Síndrome de Tietze",
            "Síndrome de Tietze esquerre sense cremor d'estómac o gastritis ( sense infusions)": "Síndrome de Tietze",
            "Pacient amb síndrome de Tietze bilateral": "Síndrome de Tietze",
            "Tendinopatia espatlla esquerra en pacient menor de 45-50 anys": "Tendinopatia espatlla",
            "Tendinopatia espatlla esquerra en pacient major de 45-50 anys": "Tendinopatia espatlla",
            "Tendinopatia espatlla esquerra calcificada": "Tendinopatia espatlla",
            "Tendinopatia espatlla esquerra amb cremors d'estómac": "Tendinopatia espatlla",
            "Cefalea general": "Cefalea",
            "Cefalea si la gola és un punt feble": "Cefalea",
            "Cefalea i dispepsia": "Cefalea",
            "Cefalea amb estrès i ansietat": "Cefalea",
            "Cefalea general amb gola dèbil i dispepsia": "Cefalea",
            "Cefalea general amb gola dèbil i ansietat": "Cefalea",
            "Cefalea general amb gola dèbil, dispepsia i ansietat": "Cefalea",
            "Cefalea general amb dispepsia i ansietat": "Cefalea",
            "Cefalea de predomini dret": "Cefalea",
            "Cefalea de predomini dret si la gola és un punt feble": "Cefalea",
            "Cefalea de predomini dret més dispèpsia": "Cefalea",
            "Cefalea de predomini dret més estrès i ansietat": "Cefalea",
            "Cefalea de predomini dret amb gola feble i dispèpsia": "Cefalea",
            "Cefalea de predomini dret amb gola feble i ansietat": "Cefalea",
            "Cefalea de predomini dret amb gola feble, dispèpsia i ansietat": "Cefalea",
            "Cefalea de predomini dret amb dispèpsia i ansietat": "Cefalea",
            "Cervicàlgia dreta": "Cervicàlgia",
            "Cervicàlgia dreta amb dispèpsia": "Cervicàlgia",
            "Torticolis aguda dreta": "Cervicàlgia",
            "Dorsàlgia interescapular dreta (tractament senzill, sense infusions)": "Dorsàlgia",
            "Dorsàlgia interescapular dreta": "Dorsàlgia",
            "Dorsàlgia dreta que arriba fins a la columna cervical (tractament senzill, sense infusions)": "Dorsàlgia",
            "Dorsàlgia dreta que arriba fins a la columna cervical": "Dorsàlgia",
            "Dorsàlgia dreta que arriba fins a la columna lumbar (tractament senzill, sense infusions)": "Dorsàlgia",
            "Dorsàlgia dreta que arriba fins a la columna lumbar": "Dorsàlgia",
            "Dorsàlgia general difusa bilateral (tractament senzill, sense infusions)": "Dorsàlgia",
            "Dorsàlgia central amb irradiació a la boca de l’estómac": "Dorsàlgia",
            "Dorsàlgia general més cervicàlgia bilateral (tractament senzill, sense infusions)": "Dorsàlgia",
            "Dorsàlgia central amb irradiació a la boca de l’estómac (tractament senzill, sense infusions)": "Dorsàlgia",
            "Migranya de predomini dret": "Migranya",
            "Migranya de predomini dret si la gola és un punt feble": "Migranya",
            "Migranya de predomini dret més dispèpsia": "Migranya",
            "Migranya de predomini dret més estrès i ansietat": "Migranya",
            "Migranya de predomini dret amb gola feble i dispèpsia": "Migranya",
            "Migranya de predomini dret amb gola feble i ansietat": "Migranya",
            "Migranya de predomini dret amb gola feble, dispèpsia i ansietat": "Migranya",
            "Migranya de predomini dret amb dispèpsia i ansietat": "Migranya",
            "Migranya general": "Migranya",
            "Migranya si la gola és un punt feble": "Migranya",
            "Migranya i dispèpsia": "Migranya",
            "Migranya amb estrès i ansietat": "Migranya",
            "Migranya general amb gola feble i dispèpsia": "Migranya",
            "Migranya general amb gola feble i ansietat": "Migranya",
            "Migranya general amb gola feble, dispèpsia i ansietat": "Migranya",
            "Migranya general amb dispèpsia i ansietat": "Migranya",
            "Neuràlgia d'Arnold de predomini dret": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret si la gola és un punt feble": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret més dispèpsia": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret més estrès i ansietat": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret amb gola feble i dispèpsia": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret amb gola feble i ansietat": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret amb gola feble, dispèpsia i ansietat": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini dret amb dispèpsia i ansietat": "Neuràigia d'Arnold",
            "Síndrome de Tietze dret (només extractes)": "Síndrome de Tietze",
            "Síndrome de Tietze dret amb ansietat": "Síndrome de Tietze",
            "Síndrome de Tietze dret amb mals digestions": "Síndrome de Tietze",
            "Síndrome de Tietze dret amb males digestions i ansietat": "Síndrome de Tietze",
            "Síndrome de Tietze bilateral": "Síndrome de Tietze",
            "Tendinopatia espatlla dreta en pacient menor de 45-50 anys": "Tendinopatia espatlla",
            "Tendinopatia espatlla dreta en pacient major de 45-50 anys": "Tendinopatia espatlla",
            "Tendinopatia espatlla dreta calcificada": "Tendinopatia espatlla",
            "Condromalàcia rotuliana dreta": "Condromalàcia rotuliana",
            "Condromalàcia rotuliana dreta amb problemes digestius (gasos, restrenyiment)": "Condromalàcia rotuliana",
            "Condromalàcia rotuliana en pacient jove (menor de 20 anys)": "Condromalàcia rotuliana",
            "Condromalàcia rotuliana en pacient jove (menor de 20 anys) amb problemes digestius (gasos i restrenyiment)": "Condromalàcia rotuliana",
            "Tendinopatia de fàscia lata dreta": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata dreta amb problemes digestius (gasos, restrenyiment)": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata en pacient jove (menor de 20 anys)": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata en pacient jove (menor de 20 anys) amb problemes digestius (gasos i restrenyiment)": "Tendinopatia de fàscia lata",
            "Trencament de fibres isquiotibial dret sense problemes digestius": "Trencament fibres isquiotibial",
            "Trencament de fibres isquiotibial dret amb problemes digestius": "Trencament fibres isquiotibial",
            "Esperó calcani dret en home": "Esperó calcani",
            "Esperó calcani dret en home amb mala circulació, varius i retenció de líquids": "Esperó calcani",
            "Esperó calcani dret en pacient home amb còlics renals previs": "Esperó calcani",
            "Esperó calcani dret en home amb hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en home amb mala circulació, varius i retenció de líquids més còlics renals previs": "Esperó calcani",
            "Esperó calcani dret en home amb mala circulació, varius i retenció de líquids més hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en home amb mala circulació, varius i retenció de líquids més còlics renals previs i hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en pacient home amb còlics renals previs i hipertensió arterial": "Esperó calcani",
            "Fascitis plantar dreta en home": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varius i retenció de líquids": "Fascitis plantar",
            "Fascitis plantar dreta en pacient home amb còlics renals previs": "Fascitis plantar",
            "Fascitis plantar dreta en home amb hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varius i retenció de líquids més còlics renals previs": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varius i retenció de líquids més hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varius i retenció de líquids més còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en pacient home amb còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Síndrome piramidal dret en home": "Síndrome piramidal",
            "Síndrome piramidal dreta en home amb prostatitis": "Síndrome piramidal",
            "Tendinopatia de la pota d’ànec homes": "Tendinopatia de pota d'oca",
            "Tendinopatia de la pota d’ànec homes amb símptomes de prostatitis": "Tendinopatia de pota d'oca",
            "Trocanteritis dreta en home": "Trocanteritis",
            "Trocanteritis dreta en home amb prostatitis": "Trocanteritis",
            "Trocanteritis bilateral en home": "Trocanteritis",
            "Artritis per microcristalls amb àcid úric alt": "Artritis per microcristalls",
            "Artritis per microcristalls amb mala circulació": "Artritis per microcristalls",
            "Artritis per microcristalls amb àcid úric alt i mala circulació": "Artritis per microcristalls",
            "Condromalàcia rotuliana esquerra (només extractes)": "Condromalàcia rotuliana",
            "Condromalàcia rotuliana esquerra": "Condromalàcia rotuliana",
            "Esperó calcani gran i evident": "Esperó calcani",
            "Esperó Calcani esquerre": "Esperó calcani",
            "Esperó calcani esquerre amb mala circulació, varius i retenció de líquids": "Esperó calcani",
            "Esperó calcani esquerre en pacient amb còlics renals previs": "Esperó calcani",
            "Esperó calcani esquerre amb hipertensió arterial": "Esperó calcani",
            "Esperó calcani esquerre amb mala circulació, varius i retenció de líquids més còlics renals previs": "Esperó calcani",
            "Esperó calcani esquerre amb mala circulació, varius i retenció de líquids més hipertensió arterial": "Esperó calcani",
            "Esperó calcani esquerre amb mala circulació, varius i retenció de líquids més còlics renals previs i hipertensió arterial": "Esperó calcani",
            "Esperó calcani esquerre en pacient amb còlics renals previs i hipertensió arterial": "Esperó calcani",
            "Esperó calcani esquerre en pacient jove (menor de 40-45 anys)": "Esperó calcani",
            "Esperó calcani, tractament base": "Esperó calcani",
            "Fascitis plantar esquerra": "Fascitis plantar",
            "Fascitis plantar esquerra amb mala circulació, varices i retenció de líquids": "Fascitis plantar",
            "Fascitis plantar esquerra en pacient amb còlics renals previs": "Fascitis plantar",
            "Fascitis plantar esquerra amb hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar esquerra amb mala circulació, varices i retenció de líquids més còlics renals previs": "Fascitis plantar",
            "Fascitis plantar esquerra amb mala circulació, varices i retenció de líquids més hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar esquerra amb mala circulació, varices i retenció de líquids més còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar esquerra en pacient amb còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar esquerra en pacient jove (menor de 40-45 anys)": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varices i retenció de líquids": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varices i retenció de líquids més còlics renals previs": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varices i retenció de líquids més hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en home amb mala circulació, varices i retenció de líquids més còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran (després de la menopausa)": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulació, varices i retenció de líquids": "Fascitis plantar",
            "Fascitis plantar dreta en pacient dona gran amb còlics renals previs": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulación, varicess i retenció de líquids més còlics renals previs": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulació, varices i retenció de líquids més hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulació, varices i retenció de líquids més còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en dona jove (abans dels 45 anys)": "Fascitis plantar",
            "Tractament senzill per a síndrome piramidal esquerre": "Síndrome piramidal",
            "Tractament per a síndrome piramidal esquerre amb infusió": "Síndrome piramidal",
            "Tractament per a síndrome piramidal esquerre (més de 45-50 anys)": "Síndrome piramidal",
            "Tendinopatia de fàscia lata esquerra en pacient jove (menor de 45 anys)": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata esquerra en pacient jove (menor de 45 anys) amb regles llargues i abundants": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata esquerra en pacient major (més de 50 anys)": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata esquerra en pacient gran amb hipertensió": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata esquerra en pacient major més ànim sota": "Tendinopatia de fàscia lata",
            "Tendinopatia de fàscia lata esquerra en pacient major amb hipertensió i ànim baix": "Tendinopatia de fàscia lata",
            "Tendinopatia de la pota d’ànec en dona": "Tendinopatia de pota d'oca",
            "Tensió ambdós trapezis amb predomini d’estrès i tensió nerviosa en pacient jove (menor de 45 anys)": "Tensió ambdós trapezis",
            "Tensió ambdós trapezis amb predomini d’estrès més suau però crònic en pacients majors (més de 45 anys)": "Tensió ambdós trapezis",
            "Trencament de fibres isquiotibial esquerre (només extractes)": "Trencament fibres isquiotibial",
            "Trencament de fibres isquiotibial esquerra en pacient cansat": "Trencament fibres isquiotibial",
            "Tractament senzill per a la trocanteritis esquerra en pacient JOVE (menor de 45 anys)": "Trocanteritis",
            "Trocanteritis esquerra en pacient JOVE (menor de 45 anys)": "Trocanteritis",
            "Trocanteritis esquerra en pacient jove (menor de 45 anys) amb regles llargues i abundants": "Trocanteritis",
            "Trocanteritis esquerra en pacient major (més de 50 anys)": "Trocanteritis",
            "Trocanteritis esquerra en pacient major amb hipertensió": "Trocanteritis",
            "Trocanteritis esquerra en pacient major més ànim sota": "Trocanteritis",
            "Trocanteritis esquerra en pacient major amb hipertensió i ànim sota": "Trocanteritis",
            "Trocanteritis bilateral en dona": "Trocanteritis",
            "Túnel carpià dret en pacient jove (menor de 45-50 anys)": "Túnel carpià",
            "Túnel carpià dret en pacient major (més de 50 anys)": "Túnel carpià",
            "Túnel carpià dret en pacient major amb hipertensió arterial (més de 50 anys)": "Túnel carpià",
            "Túnel carpià esquerre en pacient jove (menor de 45-50 anys)": "Túnel carpià",
            "Túnel carpià esquerre en pacient major (més de 50 anys)": "Túnel carpià",
            "Túnel carpià esquerre en pacient major amb hipertensió arterial (més de 50 anys)": "Túnel carpià",
            "Túnel carpià bilateral": "Túnel carpià",
            "Túnel carpià bilateral amb hipertensió arterial": "Túnel carpià",
            "Esperó calcani dret en dona gran (després de la menopausa)": "Esperó calcani",
            "Esperó calcani dret en dona gran amb mala circulació, varius i retenció de líquids": "Esperó calcani",
            "Esperó calcani dret en pacient dona gran amb còlics renals previs": "Esperó calcani",
            "Esperó calcani dret en dona gran amb hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en dona gran amb mala circulació, varius i retenció de líquids més còlics renals previs": "Esperó calcani",
            "Esperó calcani dret en dona gran amb mala circulació, varius i retenció de líquids més hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en dona gran amb mala circulació, varius i retenció de líquids més còlics renals previs i hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en dona gran amb còlics renals previs i hipertensió arterial": "Esperó calcani",
            "Esperó calcani dret en dona jove (abans dels 45 anys)": "Esperó calcani",
            "Fascitis plantar dreta en dona gran amb mala circulació, varius i retenció de líquids": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulació, varius i retenció de líquids més còlics renals previs": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulació, varius i retenció de líquids més hipertensió arterial": "Fascitis plantar",
            "Fascitis plantar dreta en dona gran amb mala circulació, varius i retenció de líquids més còlics renals previs i hipertensió arterial": "Fascitis plantar",
            "Síndrome piramidal dret en dona menor de 50 anys": "Síndrome piramidal",
            "Síndrome piramidal dret en dona amb regles doloroses": "Síndrome piramidal",
            "Síndrome piramidal dret en dona després de la menopausa": "Síndrome piramidal",
            "Trocanteritis dreta en dona menor de 50 anys": "Trocanteritis",
            "Trocanteritis dreta en dona amb regles doloroses": "Trocanteritis",
            "Trocanteritis dreta en dona després de la menopausa": "Trocanteritis",
            "Cefalea de predomini esquerre": "Cefalea",
            "Cefalea de predomini esquerre si la gola és un punt feble": "Cefalea",
            "Cefalea de predomini esquerre amb dispèpsia": "Cefalea",
            "Cefalea de predomini esquerre amb estrès i ansietat": "Cefalea",
            "Cefalea de predomini esquerre amb gola feble i dispèpsia": "Cefalea",
            "Cefalea de predomini esquerre amb gola feble i ansietat": "Cefalea",
            "Cefalea de predomini esquerre amb gola feble, dispèpsia i ansietat": "Cefalea",
            "Cefalea de predomini esquerre amb dispèpsia i ansietat": "Cefalea",
            "Cervicàlgia esquerra amb ardor d'estómac i gastritis": "Cervicàlgia",
            "Epicondilitis general": "Epicondilitis",
            "Epicondilitis i problemes digestius tipus dispèpsia (digestió lenta, gasos)": "Epicondilitis",
            "Epicondilitis amb marejos": "Epicondilitis",
            "Tractament senzill per a l'epicondilitis (sense infusions)": "Epicondilitis",
            "Tractament de base per als marejos": "Marejos",
            "Pacient amb marejos i faringitis freqüent": "Marejos",
            "Pacient amb marejos i nàusees o vòmits a la fase aguda": "Marejos",
            "Pacient amb marejos i ansietat": "Marejos",
            "Pacient amb marejos i palpitacions i/o taquicàrdies": "Marejos",
            "Pacient amb marejos , faringitis i nàusees": "Marejos",
            "Pacient amb marejos , faringitis, nàusees i ansietat": "Marejos",
            "Pacient amb marejos , faringitis , nàusees , ansietat i palpitacions": "Marejos",
            "Pacient amb marejos , nàusees i ansietat": "Marejos",
            "Pacient amb marejos , nàusees , ansietat i palpitacions": "Marejos",
            "Pacient amb marejos , nàusees i palpitacions": "Marejos",
            "Pacient amb marejos , faringitis i palpitacions": "Marejos",
            "Pacient amb marejos , faringitis i ansietat": "Marejos",
            "Cervicàlgia i marejos tractament senzill": "Marejos",
            "Migranya de predomini esquerre": "Migranya",
            "Migranya de predomini esquerre si la gola és un punt feble": "Migranya",
            "Migranya de predomini esquerre amb dispèpsia": "Migranya",
            "Migranya de predomini esquerre amb estrès i ansietat": "Migranya",
            "Migranya de predomini esquerre amb gola feble i dispèpsia": "Migranya",
            "Migranya de predomini esquerre amb gola feble i ansietat": "Migranya",
            "Migranya de predomini esquerre amb gola feble, dispèpsia i ansietat": "Migranya",
            "Migranya de predomini esquerre amb dispèpsia i ansietat": "Migranya",
            "Neuràlgia d'Arnold de predomini esquerre": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini esquerre si la gola és un punt feble": "Neuràigia d'Arnold",
            "Neuralgia d'Arnold de predomini esquerre amb dispèpsia": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini esquerre amb estrès i ansietat": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini esquerre amb gola feble i dispèpsia": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini esquerre amb gola feble i ansietat": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini esquerre amb gola feble, dispèpsia i ansietat": "Neuràigia d'Arnold",
            "Neuràlgia d'Arnold de predomini esquerre amb dispèpsia i ansietat": "Neuràigia d'Arnold",
            "Tendinopatia espatlla esquerra amb ardors d'estómac": "Tendinopatia espatlla",
            "Tensió ambdós trapezis amb predomini d'estrès i tensió nerviosa en pacient jove (menor de 45 anys )": "Tensió ambdós trapezis",
            "Tensió ambdós trapezis amb predomini d'estrès més suau però crònic en pacients majors ( més de 45 anys )": "Tensió ambdós trapezis",
            "Túnel carpià dret en pacient jove (menor de 45-50 anys )": "Túnel carpià",
            "Túnel carpià dret en pacient major ( més de 50 anys )": "Túnel carpià",
            "Túnel carpià dret en pacient major amb hipertensió arterial ( més de 50 anys )": "Túnel carpià",
        }
        
        self.zones_corpals = {
            # 🔹 Genoll
            "genoll": ["condromalàcia rotuliana", "condromalàcia rotular", "tendinitis de la cintilla iliotibial", "gonàlgia"],
            "dolor genoll": ["condromalàcia rotuliana", "condromalàcia rotular", "tendinitis de la cintilla iliotibial", "gonàlgia"],
            "genoll dret": ["condromalàcia rotular dreta", "condromalàcia rotuliana dreta", "tendinitis de la cintilla iliotibial dreta", "gonàlgia dreta"],
            "genoll esquerre": ["condromalàcia rotular esquerra", "condromalàcia rotuliana esquerra", "tendinitis de la cintilla iliotibial esquerra", "gonàlgia esquerra"],
            "dolor genoll dret": ["condromalàcia rotular dreta", "condromalàcia rotuliana dreta", "tendinitis de la cintilla iliotibial dreta", "gonàlgia dreta"],
            "dolor genoll esquerre": ["condromalàcia rotular esquerra", "condromalàcia rotuliana esquerra", "tendinitis de la cintilla iliotibial esquerra", "gonàlgia esquerra"],

            # 🔹 Gluti
            "gluti": ["síndrome del piramidal", "contractura gluti major", "punt gallet gluti", "dolor gluti"],
            "dolor gluti": ["síndrome del piramidal", "contractura gluti major", "punt gallet gluti", "dolor gluti"],
            "gluti dret": ["síndrome del piramidal dret", "contractura gluti major dreta", "punt gallet gluti dret", "dolor gluti dret"],
            "dolor gluti dret": ["síndrome del piramidal dret", "contractura gluti major dreta", "punt gallet gluti dret", "dolor gluti dret"],
            "gluti esquerre": ["síndrome del piramidal esquerre", "contractura gluti major esquerra", "punt gallet gluti esquerre", "dolor gluti esquerre"],
            "dolor gluti esquerre": ["síndrome del piramidal esquerre", "contractura gluti major esquerra", "punt gallet gluti esquerre", "dolor gluti esquerre"],

            # 🔹 Cervicals
            "cervicals": ["cervicàlgia", "hèrnia cervical", "cefalea tensional"],
            "dolor cervicals": ["cervicàlgia", "hèrnia cervical", "cefalea tensional"],
            "dolor cervical": ["cervicàlgia", "hèrnia cervical", "cefalea tensional"],
            "cervicals dretes": ["cervicàlgia dreta", "hèrnia cervical dreta", "cefalea tensional dreta"],
            "cervicals esquerres": ["cervicàlgia esquerra", "hèrnia cervical esquerra", "cefalea tensional esquerra"],
            "dolor cervicals dretes": ["cervicàlgia dreta", "hèrnia cervical dreta", "cefalea tensional dreta"],
            "dolor cervicals esquerres": ["cervicàlgia esquerra", "hèrnia cervical esquerra", "cefalea tensional esquerra"],
            "dolor cervical dret": ["cervicàlgia dreta", "hèrnia cervical dreta", "cefalea tensional dreta"],
            "dolor cervical esquerre": ["cervicàlgia esquerra", "hèrnia cervical esquerra", "cefalea tensional esquerra"],

            # 🔹 Esquena
            "esquena": ["rigidesa dorsal", "contractura lumbar", "punt gallet interescapular"],
            "dolor esquena": ["rigidesa dorsal", "contractura lumbar", "punt gallet interescapular"],

            # 🔹 Dorsal
            "dorsal": ["rigidesa dorsal", "punt gallet interescapular", "bloqueig vertebral dorsal"],
            "dolor dorsal": ["rigidesa dorsal", "punt gallet interescapular", "bloqueig vertebral dorsal"],
            "dorsal dreta": ["rigidesa dorsal dreta", "punt gallet interescapular dret", "bloqueig vertebral dorsal dret"],
            "dorsal esquerra": ["rigidesa dorsal esquerra", "punt gallet interescapular esquerre", "bloqueig vertebral dorsal esquerre"],
            "dolor dorsal dreta": ["rigidesa dorsal dreta", "punt gallet interescapular dret", "bloqueig vertebral dorsal dret"],
            "dolor dorsal esquerra": ["rigidesa dorsal esquerra", "punt gallet interescapular esquerre", "bloqueig vertebral dorsal esquerre"],

            # 🔹 Lumbar
            "lumbar": ["contractura lumbar", "punt gallet lumbar", "bloqueig vertebral lumbar"],
            "dolor lumbar": ["contractura lumbar", "punt gallet lumbar", "bloqueig vertebral lumbar"],
            "lumbar dret": ["contractura lumbar dreta", "punt gallet lumbar dret", "bloqueig vertebral lumbar dret"],
            "lumbar esquerre": ["contractura lumbar esquerra", "punt gallet lumbar esquerre", "bloqueig vertebral lumbar esquerre"],
            "dolor lumbar dret": ["contractura lumbar dreta", "punt gallet lumbar dret", "bloqueig vertebral lumbar dret"],
            "dolor lumbar esquerre": ["contractura lumbar esquerra", "punt gallet lumbar esquerre", "bloqueig vertebral lumbar esquerre"],

            # 🔹 Espatlla
            "espatlla": ["tendinopatia de l'espatlla", "lesió del manegot dels rotadors"],
            "dolor espatlla": ["tendinopatia de l'espatlla", "lesió del manegot dels rotadors", "lesió del manguito rotador"],
            "espatlla dreta": ["tendinopatia de l'espatlla dreta", "lesió del manegot dels rotadors dret", "lesión del manguito rotador dret"],
            "espatlla esquerra": ["tendinopatia de l'espatlla esquerra", "lesió del manegot dels rotadors esquerre", "lesió del manguito rotador esquerre"],
            "dolor espatlla dreta": ["tendinopatia de l'espatlla dreta", "lesió del manegot dels rotadors dret", "lesió del manguito rotador dret"],
            "dolor espatlla esquerra": ["tendinopatia de l'espatlla esquerra", "lesió del manegot dels rotadors esquerre", "lesió del manguito rotador esquerre"],

            # 🔹 Maluc
            "maluc": ["trocanteritis", "dolor sacroilíac", "bursitis trocantèrica"],
            "dolor maluc": ["trocanteritis", "dolor sacroilíac", "bursitis trocantèrica"],
            "maluc dret": ["trocanteritis dreta", "dolor sacroilíac dret", "bursitis trocantèrica dreta"],
            "maluc esquerre": ["trocanteritis esquerra", "dolor sacroilíac esquerre", "bursitis trocantèrica esquerra"],
            "dolor maluc dret": ["trocanteritis dreta", "dolor sacroilíac dret", "bursitis trocantèrica dreta"],
            "dolor maluc esquerre": ["trocanteritis esquerra", "dolor sacroilíac esquerre", "bursitis trocantèrica esquerra"],

            # 🔹 Panxa
            "panxa": ["distensió abdominal", "dispèpsia", "gastràlgia"],

            # 🔹 Cap
            "cap": ["cefalea tensional", "migranya", "neuràlgia d'Arnold"],
            "mal de cap dret": ["cefalea tensional dreta", "migranya dreta", "neuràlgia d'Arnold dreta"],
            "mal de cap esquerre": ["cefalea tensional esquerra", "migranya esquerra", "neuràlgia d'Arnold esquerra"],
            "mal de cap": ["cefalea tensional", "migranya", "neuràlgia d'Arnold"],

            # 🔹 Pit
            "pit": ["dolor toràcic", "punxades al pit", "síndrome de Tietze"],
            "dolor pit": ["dolor toràcic", "punxades al pit", "síndrome de Tietze"],
            "pit dret": ["dolor toràcic dret", "punxades al pit dret", "síndrome de Tietze dret"],
            "pit esquerre": ["dolor toràcic esquerre", "punxades al pit esquerre", "síndrome de Tietze esquerre"],
            "dolor pectoral": ["dolor toràcic", "punxades al pit", "síndrome de Tietze"],
            "dolor pectoral dret": ["dolor toràcic dret", "punxades al pit dret", "síndrome de Tietze dret"],
            "dolor pectoral esquerre": ["dolor toràcic esquerre", "punxades al pit esquerre", "síndrome de Tietze esquerre"],

            # 🔹 Canell
            "canell": ["síndrome del túnel carpià", "tendinitis del canell", "artrosi del canell"],
            "dolor canell": ["síndrome del túnel carpià", "tendinitis del canell", "artrosi del canell"],
            "canell dret": ["síndrome del túnel carpià dret", "tendinitis del canell dret", "artrosi del canell dret"],
            "canell esquerre": ["síndrome del túnel carpià esquerre", "tendinitis del canell esquerre", "artrosi del canell esquerre"],
            "dolor canell dret": ["síndrome del túnel carpià dret", "tendinitis del canell dret", "artrosi del canell dret"],
            "dolor canell esquerre": ["síndrome del túnel carpià esquerre", "tendinitis del canell esquerre", "artrosi del canell esquerre"],

            # 🔹 Turmell
            "turmell": ["esquinç de turmell", "tendinitis d'Aquil·les", "bursitis retrocalcània"],
            "dolor turmell": ["esquinç de turmell", "tendinitis d'Aquil·les", "bursitis retrocalcània"],
            "turmell dret": ["esquinç de turmell dret", "tendinitis d'Aquil·les dreta", "bursitis retrocalcània dreta"],
            "turmell esquerre": ["esquinç de turmell esquerre", "tendinitis d'Aquil·les esquerra", "bursitis retrocalcània esquerra"],
            "dolor turmell dret": ["esquinç de turmell dret", "tendinitis d'Aquil·les dreta", "bursitis retrocalcània dreta"],
            "dolor turmell esquerre": ["esquinç de turmell esquerre", "tendinitis d'Aquil·les esquerra", "bursitis retrocalcània esquerra"],

            # 🔹 Peu
            "peu": ["fascitis plantar", "metatarsàlgia", "espoló calcani", "esperó calcani"],
            "dolor peu": ["fascitis plantar", "metatarsàlgia", "espoló calcani", "esperó calcani"],
            "peu dret": ["fascitis plantar dreta", "metatarsàlgia dreta", "espoló calcani dret", "esperó calcani dret"],
            "peu esquerre": ["fascitis plantar esquerra", "metatarsàlgia esquerra", "espoló calcani esquerre", "esperó calcani esquerre"],
            "dolor peu dret": ["fascitis plantar dreta", "metatarsàlgia dreta", "espoló calcani dret", "esperó calcani dret"],
            "dolor peu esquerre": ["fascitis plantar esquerra", "metatarsàlgia esquerra", "espoló calcani esquerre", "esperó calcani esquerre"],
        }

        self.vísceres = {
            "fetge": {
                "símptomes": [
                    "Cervicàlgia dreta (incloses les produïdes per torticoli, hèrnia discal o punts gallet).",
                    "Dolor interescapular dret",
                    "Tendinopatia de l'espatlla dreta",
                    "Cefalea tensional amb tendència a la migranya al costat dret predominantment (però també pot ser esquerra). El mal de cap pot arribar fins a l'ull.",
                    "Cefalea per gana és símptoma clar de mal funcionament de fetge",
                    "Dolor a la zona suboccipital dreta",
                    "Neuràlgia d'Arnold",
                    "Contractures i punts gallet interescapulars drets.",
                    "Rigidesa dorsal, sobretot T5 i T6 (síndrome de Tietze dret)",
                    "Tendinopaties de l'espatlla dreta",
                    "Dolor a la graella costal dreta, a la zona anterior i lateral (normalment indica que el fetge està congestionat, se soluciona amb dieta i depuració amb herbes).",
                    "Dorsal rectificada",
                    "Cefalea post estrès",
                    "mal de cap dret", 
                    "cefalea tensional dreta", 
                    "migranya dreta",
                    "dolor espatlla dreta",
                    "dolor interescapular dret",
                    "dolor canell",
                    "síndrome túnel carpià"
                ],
                "altres_símptomes": [
                    "Cansament al matí acabat de despertar.",
                    "Icterícia.",
                    "Pell greixosa, acné i problemes de pell que provoquen coïssor i picor.",
                    "Quistos sebacis a parpelles, cuir cabellut i en qualsevol zona.",
                    "Dislipidèmies (colesterol i triglicèrids alts).",
                    "Desig de menjar xocolata, formatge, cafè, embotits, fregits.",
                    "Trastorns del son: somni inquiet durant les primeres hores de la nit i somnolència excessiva després de dinar.",
                    "Problemes digestius, cremor d'estómac, reflux i/o amargor a la boca.",
                    "De forma secundària, certs tipus de restrenyiment poden estar lligats a la disfunció hepàtica.",
                    "Alè que desprèn olor a poma o acetona.",
                    "Femtes més blanques (per falta de bilis).",
                    "Faringitis o amigdalitis repetitives."
                ],
                "origen": {
                    "caràcter": "Impetuós, valent, amant de la novetat, enèrgic, entusiasta, però que cau fàcilment en el descontrol, l'estrès i, en casos extrems, en agressivitat.",
                    "emocions": [
                        "Enfadament",
                        "Ràbia per alguna cosa que no acceptem o perquè no ens sentim acceptats",
                        "Sobrecàrrega de responsabilitats",
                        "Estrès laboral",
                        "Estrès de vida"
    ],
                    "nutrició": "Lactis, taronja i mandarina, greixos saturats, dolços manufacturats, cacau, cafè, carn processada, alcohol, medicaments hepatotòxics com els analgèsics i antiinflamatoris no esteroïdals (AINES)."
            },
                "exploració": [
                    "Espatlla dreta antepulsada i més baixa que l'esquerra.",
                    "Dorsal rectificada (fer-li flexionar l'esquena per veure-ho, tot i que normalment es veu estant erguit).",
                    "Dolor a la palpació T5.",
                    "Punt gallet angle escàpula dreta.",
                    "Rigidesa dorsal."
    ],
                "a_tenir_en_compte": [
                    "Espatlla dreta avançada i més baixa que l'esquerra té a veure amb estrès.",
                    "Dolor graella costal dreta és per sobrecàrrega del fetge. Amb dieta i depuració del fetge marxa en 4-5 dies.",
                    "El 90% de les migranyes són per histaminosi. Si es redueix al màxim els aliments que porten histamina, se'n van. El glutamat monosòdic que porten molts processats, és alt en histamina.",
                    "Sempre que és fetge o vesícula amb repercussió a la zona cervical cal receptar plantes per al sistema nerviós.",
                    "Condicions com la diabetes, hipotiroidisme o embaràs poden augmentar la predispisició a la síndrome del túnel carpià, ja que generen retenció de líquids i canvis inflamatoris en el cos."
                ]
            },
            "vesícula biliar": {
                "símptomes": [
                    "Cefalea tensional, de vegades només a l'esquerra, amb mareig-inestabilitat.",
                    "Dolor clatell esquerre, de vegades acompanyat de mareig-inestabilitat.",
                    "Cervicàlgia esquerra, hèrnia discal esquerra (sobretot si el dolor arriba al clatell i amb marejos-inestabilitat).",
                    "Dolor suboccipital esquerre que pot irradiar al fons de l'ull",
                    "Tortícoli esquerra repetitiva",
                    "Nàusees, vòmits",
                    "Mareig, inestabilitat, espessor mental, i si augmenta en un centre comercial, gairebé segur que és de vesícula biliar",
                    "Sensació de pesadesa-tensió en ambdós trapezis.",
                    "Tendinopatia espatlla esquerra associada a inestabilitat cervical.",
                    "Rigidesa dorsal sobretot T4 i T5.",
                    "mal de cap esquerre", 
                    "cefalea tensional esquerra", 
                    "migranya esquerra",
                    "dolor esparlla esquerra",
                    "dolor colze esquerre",
                    "dolor canell",
                    "síndrome túnel carpià",
                    "epicondilitis"
                ],
                "altres_símptomes": [
                    "Sensació de nàusees o fàstics.",
                    "Nàusees matutines davant de diferents olors o menjars.",
                    "Nàusees davant de situacions emocionals difícils.",
                    "Hipersensibilitat olfactiva i visual.",
                    "Sinusitis i rinitis (infeccions sinus nasals).",
                    "Estrenyiment per manca de peristaltisme.",
                    "Dispèpsia funcional: digestió lenta per mal alliberament de secrecions biliars.",
                    "Sensació de plenitud i inflor després d’ingerir aliments.",
                    "Formació de càlculs o pedres a la vesícula biliar.",
                    "Risc d'obstrucció de les vies biliars i còlics biliars.",
                    "Pancreatitis derivada de càlculs biliars."
                ],

                "origen": {
                    "caràcter": "Introvertit, amb tendència al perfeccionisme i excessiva necessitat de control i autocrítica. Falta de flexibilitat que fa que els costi acceptar situacions noves. Viu en alerta i tensió tot el dia perquè necessita tenir-ho tot controlat.",
                    "emocions": [
                        "Amargor, disgust, impotència, frustració, dificultat a digerir esdeveniments nous. Al final és una ràbia per una situació no digerida que es transforma en amargor.",
                        "Estrès"
                    ],
                    "nutrició": "Igual que el fetge, tot i que en la vesícula biliar no és tan important el factor nutrició. Però molt important el consum de cafè. La resta, menys important: lactis, taronja i mandarina, greixos saturats, dolços manufacturats, cacau, carn processada, alcohol, medicaments hepatotòxics com els analgèsics i antiinflamatoris no esteroïdals (AINES)."
                },
                "exploració": [
                    "Espatlla esquerra més elevada i enrere respecte de la dreta.",
                    "Tortícoli esquerra repetitiva.",
                    "Dolor a la palpació de T4 i T5.",
                    "Dolor a la palpació de la vora medial de l'escàpula esquerra.",
                    "Punt gallet suboccipital esquerre.",
                    "Dolor cervicodorsal esquerre."
                ],
                "a_tenir_en_compte": [
                    "Si l'espatlla esquerra està més elevada i enrere respecte de la dreta, té a veure amb estrès.",
                    "Si hi ha dolor cervical per vesícula, acostuma a haver-hi una hipersensibilitat olfactiva i visual.",
                    "Mareig-inestabilitat pot ser causat per vesícula biliar.",
                    "Sempre que és fetge o vesícula amb repercussió a la zona cervical cal receptar plantes per al sistema nerviós.",
                    "Condicions com la diabetes, hipotiroidisme o embaràs poden augmentar la predispisició a la síndrome del túnel carpià, ja que generen retenció de líquids i canvis inflamatoris en el cos."
                ]
            },    
            "estómac": {
                "símptomes": [
                    "Torticoli i hèrnia discal cervical esquerra per inestabilitat. Contractura i punt gallet interescapular esquerre.",
                    "Sensació de pesadesa braç esquerre, formigueig",
                    "Bloqueig vertebral dorsal (sobretot T6, o T5-T6 si ve de fetge).",
                    "Bloqueig cervical C7 banda esquerra",
                    "Dolor just sota la darrera costella esquerra",
                    "Dolor muscular-costal i cartílag costoesternal esquerre, sobretot costelles 5, 6 i 7 (síndrome de Tietze).",
                    "Punts gallet a l'infraespinós esquerre.",
                    "Tendinopatia espatlla esquerra per antepulsió de l'espatlla o facilitació segmentària cervical per hipermobilitat.",
                    "dolor pectoral esquerre",
                    "dolor espatlla esquerra",
                    "dolor interescapular esquerre"
                ],
                "altres_símptomes": [
                    "Acidesa, reflux, cremor d'estómac o gastritis crònica.",
                    "Sensació d'inflor immediatament després d’ingerir aliments.",
                    "Ronqueres i tos seca en tombar-se després de sopar.",
                    "Molèstia a la zona infracostal esquerra, típica de la gastritis.",
                    "Els símptomes poden aparèixer l’endemà d’un àpat copiós."
],
                "origen": {
                    "caràcter": "Obsessiu amb les regles, amb el treball, amb la responsabilitat.",
                    "emocions": [
                        "Excés de responsabilitat fonamentalment a nivell material: excés de responsabilitat a la feina o els estudis, per exemple, i dificultat de desconnectar.",
                        "Estar pensant tot el dia en el que s'ha d'organitzar, no desconnectar d'això (per exemple, tenir obres a casa i obsessionar-se amb el procés, tenir pressa perquè s'acabi i no poder desconnectar).",
                        "Emocions de fetge: enfadament, ràbia per alguna cosa que no acceptem. O perquè no ens sentim acceptats.",
                        "Estrès."
                    ],
                    "nutrició": "Igual que el fetge, ja que es veu directament afectat pel mal funcionament del fetge, i eliminar també alls, picant i vinagre."
},
                "exploració": [
                    "Espatlla esquerra més elevada i enrere respecte de la dreta.",
                    "Dolor a la palpació de T6 (o T5-T6 si ve de fetge).",
                    "Punt gallet infraespinós esquerre.",
                    "Dolor a la vora medial de l'escàpula esquerra.",
                    "Dolor a la palpació sota la darrera costella esquerra.",
                    "Dolor a la palpació del cartílag costal esquerre (costelles 5, 6 i 7) (Síndrome de Tietze).",
                    "Dolor a la palpació de l'epigastri (fort dolor a la boca de l'estómac)."
                ],
                "a_tenir_en_compte": [
                    "Si l'espatlla esquerra està més elevada i enrere respecte de la dreta, té a veure amb estrès.",
                    "Si hi ha dolor cervical per estómac, acostuma a haver-hi una hipersensibilitat olfactiva i visual.",
                    "Una emoció típica en el problema digestiu és que costa deixar anar les coses.",
                    "Sempre que és estómac o intestí prim, cal receptar plantes per al sistema nerviós."
                ]
            },
          
            "sistema circulatori i cor": {
                "símptomes": [
                    "Dolor del trapezi esquerre fins a la clavícula.",
                    "Sensació de pesadesa a ambdós trapezis, com si portessis una motxilla carregada.",
                    "Bloqueig dorsal T2 i T3 (sensació lleugera de torticoli)",
                    "Inflamació condro-costal esquerra. 1a, 2a i 3a costella",
                    "Dolor axil·lar que de vegades irradia a la zona interna del braç fins al colze esquerre",
                    "dolor trapezi esquerre",
                    "dolor axil·lar esquerre",
                    "dolor aixella esquerra"
                ],
                "altres_símptomes": [
                    "Palpitacions (sensació que el cor batega amb violència).",
                    "Taquicàrdies funcionals: alteració súbita del ritme cardíac en moments de repòs.",
                    "Sensació opressiva al pit i ansietat.",
                    "No hi ha cap símptoma en fer exercicis físics extrems."
                ],
                "origen": {
                    "caràcter": "Persona nerviosa amb tendència a l'ansietat.",
                    "emocions": [
                        "Ansietat provocada per conflictes quant a la percepció de l'amor: amor malentès, sensació de no ser estimat incondicionalment pels pares.",
                        "Narcissisme emocional: en el fons ve de la percepció de falta d'amor incondicional dels pares, perquè aquest narcissisme és intentar fer sempre el correcte davant els altres, buscar permanentment la seva aprovació.",
                        "Ansietat per una cosa positiva: unes vacances que estem preparant per exemple, estic organitzant el meu casament...",
                        "Influència indirecta de les emocions que afecten la vesícula biliar com l'amargor i el disgust i l'estrès viscut de forma introvertida."
                    ],
                    "nutrició": "Normalment la dieta acidificant, que contribueix a l'estrès. Consum d'excitants com el cafè. Molt consum de carn vermella i carn processada, molta proteïna animal."
                },
                
                "exploració": [
                    "Dolor a la T7"
                ],
                "a_tenir_en_compte": []             

            },
            "intestí prim": {
                "símptomes": [
                    "Lumbàlgies amb mal de predomini dret per bloqueig de la L4 a FRSi (flexió-inclinació i rotació esquerra).",
                    "Sacroileitis dreta per obertura ilíaca.",
                    "Síndrome piramidal dreta (de debò, no fals).",
                    "Dolor engonal dreta per inflamació del psoes.",
                    "Ciàtica dreta per hèrnia discal L4-L5.",
                    "Trencament de fibres repetitives de bíceps femoral dret.",
                    "Condromalàcia rotuliana dreta, a causa d’un lleu desplaçament de la ròtula cap a fora perquè el to del vast extern està augmentat.",
                    "Tendinitis de la cintilla iliotibial dreta o genoll del corredor.",
                    "Bloqueig del cap del peroné dret.",
                    "Derivat de l’augment de to de la cadena d’obertura dreta, pot haver-hi una fascitis plantar dreta.",
                    "Dolor engonal dret",
                    "dolor gluti dret",
                    "dolor genoll dret"
                ],
                "altres_símptomes": [
                    "Qualsevol alteració del procés digestiu: diarrees, restrenyiment, gasos, digestió lenta.",
                    "Distensió abdominal i sensació de panxa plena després de menjar.",
                    "Molta son després de menjar, fins i tot si s’ha menjat poc.",
                    "Sensació desmesurada d'opressió per la roba ajustada o el cinturó.",
                    "Bruxisme (apretar les dents durant la nit).",
                    "Irritació ganglis gola.",
                    "Capa groguenca a la zona posterior de la llengua.",
                    "Irritació de la llengua, llengua pastosa.",
                    "Aparició d'acne o vermellors al front.",
                    "Símptomes celiaquia: distensió abdominal, gasos, dispèpsia, reflux, intolerància a altres aliments."
                 ],
                "origen": {
                    "caràcter": "Persona nerviosa amb tendència a la preocupació.",
                    "emocions": [
                        "Preocupació, ansietat pel que ha de venir i excés d'anàlisi dels problemes.",
                        "També es pot veure malmès per les emocions del fetge, com l'estrès."
                    ],
                    "nutrició": "Excessiu consum de sucre i aliments refinats per eliminar la fibra com els cereals i l'arròs blanc. També cal posar atenció en els aliments que afecten el fetge, ja que de forma indirecta afectaran el funcionament del budell prim. Ús abusiu d'antibiòtics, que afecten negativament la microbiota intestinal."
                },
                "exploració": [
                    "En bipedestació, cresta ilíaca dreta lleugerament més alta i giba lumbar esquerra.",
                    "En bipedestació, EIAS esquerra més alta que la dreta (certa rotació lumbar esquerra, tronc orientat lleugerament a l’esquerra).",
                    "Assegut, cresta ilíaca dreta igualment més alta que l'esquerra.",
                    "En el test de flexió assegut de les EIPS, més rang de mobilitat de la banda dreta (cal posar el dit just sota les EIPS per poder seguir bé el moviment).",
                    "Tombat en supí, cama dreta sensiblement més llarga que l’esquerra (falsa cama curta). (Per fer el test bé, abans apretar els genolls fins al tronc perquè no ens enganyi la postura).",
                    "En supí, augment rotació externa cama dreta.",
                    "A la camilla en supí amb els genolls doblegats a 90 graus (angle dels malucs a 90 graus respecte al tronc) el genoll dret queda clarament més a baix (per fer el test bé, abans apretar els genolls fins al tronc perquè no ens enganyi la postura). Això és molt simptomàtic d’intestí prim.",
                    "Sensibilitat especial a la palpació L4-L5 banda dreta. Aquest símptoma és pràcticament imprescindible en cas d’intestí prim.",
                    "En prono, demanar-li extensió tronc, i s’aixeca més la lumbar esquerra que la dreta.",
                    "En supí, per acabar de confirmar, li ha de fer mal la fosa ilíaca dreta, molt més que l’esquerra.",
                    "Irritació dolorosa T12."
                ],
                "a_tenir_en_compte": [
                    "Principal causa: mescla de mala alimentació + estrès.",
                    "L’intestí prim irritat augmenta el to de la cadena d’apertura dreta: glutis, piramidal, quadrat crural, obturadors, tibials, bessó intern, bíceps llarg, bíceps curt, extensor llarg primer dit peu, tensor fàscia lata.",
                    "Si no hi ha dolor L4-L5 dret + dolor fosa ilíaca dreta + dolor T12, pràcticament segur que no ve d’intestí prim. En canvi, si es compleixen les 3 condicions, pràcticament segur que ve d'intestí prim.",
                    "La sensibilitat al gluten en nens i adolescents dóna pocs símptomes digestius i els primers avisos són en el sistema múscul-esquelètic: dolor a l’engonal dreta, genoll del corredor, tendinitis fàscia lata dreta, dolor rotulià dret... En menys ocasions, dolor lumbar.",
                    "En adults, la sensibilitat al gluten dóna els mateixos símptomes que en nens, més: afectació lumbar i a la llarga hèrnia discal lumbar dreta entre L4-L5. Només amb aquest darrer símptoma ja ens ha de fer pensar en problemes intestinals.",
                    "A l’exploració: rigidesa T12, rigidesa lumbosacra dreta, inestabilitat lumbar."
                ]
            },
            "pàncrees": {
                "símptomes": [
                    "Bloqueig vertebral dorsal a T7 i T8",
                    "Contractura interescapular esquerra a l'alçada de T7-T8-T9",
                    "De vegades irradiació costal dolorosa",
                    "Dolor interescapular esquerre"
                ],
                "altres_símptomes": [
                    "Alteració o inestabilitat en els nivells de glucosa a la sang.",
                    "Atacs de fam incontrolable a mig matí o mitja tarda.",
                    "Somnolència en horari d'11 a 12 AM o de 5 a 6 PM.",
                    "Infeccions per fongs (a qualsevol part del cos, sobretot a la boca).",
                    "Disminució de la capacitat immunitària."
                ],
                "origen": {
                    "caràcter": "Obsessió i dificultat en desconnectar de les responsabilitats.",
                    "emocions": [
                        "Obsessió en l'àmbit familiar que es tradueix en la dificultat de desconnectar dels problemes i responsabilitats familiars.",
                        "Estrès que ve de la vesícula biliar."
                    ],
                    "nutrició": "Massa hidrats de carboni d'absorció ràpida."
                },
                "exploració": [
                        "Dolor T4-T5 (perquè moltes vegades ve de problemes derivats de la disfunció de vesícula biliar)."
                    ],
                "a_tenir_en_compte": []
            },
            "colon irritable": {
                "símptomes": [
                    "En algunes ocasions, bloqueig de L1 i L2.",
                    "Lumbàlgia a la zona lumbar baixa, de predomini esquerre, moltes vegades cronificada durant anys.",
                    "Ciàtica esquerra per hèrnia discal L5-S1.",
                    "Dolor abdominal esquerre per espasme de còlon descendent.",
                    "Dolor engonal esquerre per tendinitis del psoes."
                    "dolor abdominal esquerre",
                    "dolor engonal esquerre",
                    "dolor lumbar",
                    "dolor gluti esquerre"
                ],
                "altres_símptomes": [
                            "Alternança entre diarrees davant de qualsevol estrès i posterior restrenyiment.",
                            "Diarrees nervioses en situacions com exàmens, viatges, situacions d'estrès."
                 ],
                "origen": {
                    "caràcter": "Individu nerviós.",
                    "emocions": [
                        "Individus als quals els costa dir no a les demandes dels altres, perquè se senten esclaus de donar sempre una bona imatge i d'aquesta manera, sentir-se valorats pels altres.",
                        "En el fons és por a no ser acceptats pels altres.",
                        "Ve d'un patró de la infància: no tenir l'amor incondicional dels pares."
                    ],
                    "nutrició": "No és mai la causa perquè el problema és d'origen nerviós."
                },
                "exploració": [
                    "Dismetria maluc. Cresta ilíaca esquerra més baixa respecte a la dreta.",
                    "Assegut: dolor i mobilitat excessiva cresta ilíaca esquerra amb el test de flexió assegut.",
                    "En supí, test de flexió de genolls, genoll esquerre més alt. I cama esquerra curta rotada internament.",
                    "En prono, giba esquerra i sacre orientat a l’esquerra.",
                    "Dolor L1 lateral esquerra (aquesta vèrtebra normalment no fa mal en cap altre cas)."
                ],
                "a_tenir_en_compte": [
                    "Com que ve d’un patró de la infància (amor condicionat dels pares), costa de revertir.",
                    "El colon irritable és l'única somatització de colon (en estadis més greus es transforma en colitis ulcerosa i malaltia de Crohn, però ja no se somatitza músculoesquelèticament).",
                    "Els símptomes músculoesquelètics, que produeixen un tancament ilíac esquerre, tenen molt a veure amb la sobretensió de la cadena de tancament esquerra (adductors i psoas ilíac especialment), de manera que s'ha de relaxar aquesta cadena.",
                    "L'exploració és la mateixa que la d'intestí prim. L'única diferència és el test de flexió, que dóna hipermobilitat a la dreta en el cas de l'intestí prim."
                ]
            },
            
            "ronyó": {
                "símptomes": [
                    "Rigidesa i bloqueig de T9-T10-T11 amb dolor lumbar alt (uni o bilateral) que en ocasions pot irradiar-se cap a l'abdomen.",
                    "Rigidesa lumbosacra amb dolor lumbar baix de predomini esquerre.",
                    "Trocanteritis esquerra.",
                    "Sacroileitis esquerra.",
                    "Síndrome del piramidal esquerre, que en el fons és un fals síndrome del piramidal.",
                    "Ciàtica esquerra per hèrnia discal L5-S1.",
                    "Dolor a l'engonal per tendinitis del psoes.",
                    "Lesions del bíceps femoral esquerre.",
                    "Tendinitis de la cintilla iliotibial esquerra o tendinitis del corredor.",
                    "Condromalàcia rotuliana esquerra o síndrome d'hiperpressió rotuliana.",
                    "Tendinitis taló d'Aquil·les esquerre.",
                    "Bloqueig de cap de peroné esquerre.",
                    "Esperó calcani a qualsevol dels dos peus (vé igualment de l'hèrnia, però amb un problema afegit de ronyó: els microcristalls).",
                    "Tendinitis per microcristalls als extensors del turmell, a qualsevol turmell.",
                    "Lumbàlgia post estrès.",
                    "Fascitis plantar esquerra.",
                    "Artritis coxofemoral esquerra.",
                    "dolor lumbar alt bilateral",
                    "Dolor lumbar esquerre",
                    "dolor ambdós turmells",
                    "Dolor ambdós peus",
                    "dolor peus",
                    "dolor canell"
                ],
                "altres_símptomes": [
                    "Cansament, manca d'energia durant el dia.",
                    "Decaïment, apatia que pot evolucionar a sentiments de depressió.",
                    "Retenció de líquids, sobretot a les cames.",
                    "Hipertensió arterial que indicaria un signe de patologia renal.",
                    "Malsons recurrents durant la nit.",
                    "Sensació de son superficial, no reparadora.",
                    "Disminució de la menstruació, fins i tot amenorrea.",
                    "Elevació de la creatinina i àcid úric en anàlisis de sang.",
                    "Constant sensació de fred principalment de cintura cap avall.",
                    "Pèrdua de memòria o dificultat en l’esforç intel·lectual."
                ],
                "origen": {
                    "caràcter": "Persona introvertida, prudent, inadvertida, observadora i amb tendència a l'aïllament.",
                    "emocions": [
                        "Por. Manca d'autoestima, inseguretat, sentiment d'abandó.",
                        "També pot ser una persona que rep una pressió molt gran i se sent desbordada per la responsabilitat perquè té por de no poder estar a l’altura.",
                        "Aferrament al passat, malenconia... emocions d’intestí gros, però que sostingudes acaben afectant l’element següent i evolucionen cap a tristor.",
                        "Estrès mantingut que després de passar per vesícula biliar/fetge i sistema nerviós acaba rebotant en l’element anterior, el ronyó."
                    ],
                    "nutrició": "Marisc, lactis (especialment el formatge), carn processada com els embotits, aliments amb sal afegida (atenció als encurtits), cafè, alcohol, dolços, verdures amb àcid oxàlic i medicaments nefrotòxics d'ús freqüent, com els antiinflamatoris."
                },
                "exploració": [
                    "Pèrdua de lordosi lumbar a causa de la retroversió pèlvica (com els animals, com si poséssim la cua entre les cames per por).",
                    "Centre de gravetat anterior, individu rectilini?",
                    "En bipedestació la cresta ilíaca esquerra estarà més alta, assegut també.","En el test de flexió notarem més rigidesa al voltant de l’EIPS esquerra.",
                    "En decúbit supí la cama esquerra curta i rotada externament.",
                    "Palpem les dues carilles laterals L5-S1, les anem pressionant alternant la pressió a una i a l’altra per irritar-les, i si és ronyó l’esquerra acabarà fent més mal. I la notarem més rígida també.",
                    "Palpant T9-T10-T11, alguna de les tres ha de fer mal.",
                    "Dolor lumbar generalitzat, ambdues bandes, gairebé segur és de ronyó.",
                    "El més típic és que la disfunció de ronyó vingui d’esgotament per estrès, llavors segurament li fa mal també la T4 o T5. Comprovar per saber què treballar."
                ],
                "a_tenir_en_compte": [
                    "La trocanteritis esquerra és gairebé segur símptoma de ronyó (indica esgotament que molt possiblement ve d'estrès sostingut durant anys més mala alimentació, més antiinflamatoris).",
                    "La síndrome del piramidal esquerre també és molt probable que vingui de ronyó.",
                    "Si l'origen és l'estrès, cal tractar també vesícula biliar i fetge.",
                    "Tots els símptomes pèlvics preferentment es donen en gent que té el centre de gravetat anterior (rectilinis).",
                    "Si la persona et diu que des de jove li ha fet sempre més mal la zona lumbar, segurament és por o inseguretat.",
                    "En canvi, si diu que de jove el que li feia mal era l’esquena alta o cervicals, serà esgotament per estrès, que primer ha començat per vesícula biliar, s’ha traslladat al sistema nerviós i ha acabat al ronyó.",
                    "Quan veus una ciàtica o qualsevol problema músculoesquelètic que ve de ronyó en una persona de menys de 55 anys el més probable és que hi estigui implicat també el sistema nerviós i la vesícula i el fetge. Per tant, afegir herbes d'aquests dos sistemes.",
                    "Dolor lumbar alt sord però molt molest podria indicar un inici de pironefritis (còlic renal). Per comprovar-ho, donar-li un copet al ronyó, si fa mal o et deixa endolorit és que és pironefritis. Amb dieta i herbes se soluciona molt bé.",
                    "Una ciàtica esquerra sobtada, a vegades en una persona jove, i sense cap antecedent anterior de traumatologia, indica normalment que la persona se sent desbordada per una gran responsabilitat que li genera molta pressió. En el fons és por a no estar a l'altura de la responsabilitat."
                ]
            },
            "bufeta": {
                "símptomes": [
                    "Dolor lumbar a l'alçada de L3.",
                    "Dolor-pressió damunt del pubis.",
                    "Dolor-tivantor-tensió dels dos bessons, incoherent amb l’activitat física que fa el o la pacient.",
                    "Dolor a les dues tíbies (menys freqüent).",
                    "Trencament de fibres del bessó.",
                    "Tendinitis dels tendons d'Aquil·les.",
                    "dolor bessons",
                    "dolor tíbies",
                    "dolor tendons d'Aquil·les"
                ],
                "altres_símptomes": [
                    "Infeccions urinàries recurrents, desencadenades per fred i humitat.",
                    "Febre, ardor o coïssor en orinar, pressió o dolor al pubis.",
                    "Orinar moltes vegades però poca quantitat.",
                    "Orinar sang.",
                    "Orinar amb freqüència i poca quantitat.",
                    "Dolor suprapúbic irradiat a zona lumbar."
                ],
                "origen": {
                    "caràcter": "Persona reservada amb tendència a tenir vergonya. Prudent, insegura, persona que passa en segon pla, apagada i amb baixa autoestima.",
                    "emocions": [
                        "Culpa en general. Sensació de culpa per disfrutar, perquè hi ha una culpa de base, o culpa perquè la meva mare em fa sentir fatal quan la deixo sola, etc.",
                        "Dones: culpa relacionada amb les relacions sexuals.",
                        "Desesperança: veure el futur negre, algun procés que se’t fa molt dur i no li veus sortida.",
                        "Veure el teu terreny trepitjat, no veure’s tingut en compte, més en dones. No imposar-se en una situació de parella nova, etc. Té a veure amb la inseguretat.",
                        "Estrès que ve de fetge i que també implica el sistema nerviós."
                    ],
                    "nutrició": "En la bufeta la nutrició mai és el més important."
                },
                "exploració": [],  # No hi ha informació en aquest apartat segons el Word
                "a_tenir_en_compte": [
                    "Dolor-tivantor-tensió dels dos bessons, incoherent amb l’activitat física que fa el o la pacient, és segur símptoma de bufeta.",
                    "Si una persona que no fa activitat física forta té un trencament de fibres amb un gest normal, o pujant escales, etc., gairebé segur que és bufeta. I si un trencament de fibres de bessons després de tres setmanes continua fent mal, segurament també és bufeta. Més freqüent en dones.",
                    "Quan hi ha trencament de fibres normalment hi està implicada una emoció de fetge (ràbia). Moltes vegades també faringitis repetitives (ràbia no expressada per una situació no acceptada)."
                    ]
            },
            
            "sistema ginecològic femení": {
                "símptomes": [
                    "Bloqueig de L5 a ERSd (extensió, inclinació i rotació dreta)",
                    "Dolor lumbar de predomini dret",
                    "Hèrnia discal L5-S1 dreta",
                    "Dolor baix ventre",
                    "Dolor a les dues EIAS",
                    "Sacroileitis dreta",
                    "Trocanteritis dreta (secundari al bloqueig de L5 a ERSd)",
                    "Síndrome del piramidal dret (secundari al bloqueig de L5 a ERSd)",
                    "Dolor sota ventre (en el cas de l'úter)",
                    "Dolor engonal dreta per tendinitis del psoes-ilíac",
                    "Ciàtica dreta per lesió discal L5-S1",
                    "Tendinitis pata de ganso, sobretot del costat dret",
                    "Dolor i inflamació sota els mal·lèols externs i sinus del tars",
                    "Dolor al canell",
                    "Condromalàcia rotuliana ambdues cames (especialment noies joves i amb regles llargues i abundants)",
                    "Esperó calcani dret",
                    "dolor turmells",
                    "dolor peus",
                    "dolor lumbar",
                    "dolor genoll dret",
                    "dolor gluti dret"
                ],
                "altres_símptomes": [
                    "Regla dolorosa, amb dolor lumbar.",
                    "Cicles irregulars.",
                    "Regles molt abundants i de més de 5 dies.",
                    "Regles molt escasses i de curta durada.",
                    "Acné a la barbeta i al voltant de la boca.",
                    "Alteració de la coloració del sagnat a les regles."
                  ],
                "origen": {
                    "caràcter": "Persona introvertida, prudent, insegura, inadvertida, observadora, amb baixa autoestima i amb tendència a l'aïllament.",
                    "emocions": [
                        "Causa més comú: esgotament com a conseqüència de l’estrès, post estrès.",
                        "Autoestima danyada en l’àmbit de la parella, per sentir-se enganyada, abandonada, traïda, o poc valorada, no sentir-se una prioritat per a la parella.",
                        "Pors amb els fills, molt freqüents en dones (el problema és que tenen por elles i les traslladen als fills).",
                        "Tristor per malalties de llarga evolució amb deteriorament dels pares, habitual en gent que cuida dels pares.",
                        "Sensació de pèrdua de vàlua com a dona -que en el fons és pèrdua d’autoestima- com a conseqüència de la menopausa.",
                        "També conflictes amb la maternitat o no maternitat (no poder tenir fills perquè t’has despertat massa tard, o haver tingut fills i adonar-se que era per pressió social i no era perquè ella volia, etc.).",
                        "O també sentir-se degradada en el treball, poc valorada (això sobretot passa en dones que no tenen fills, perquè bolquen tota la càrrega emocional en la feina, és la seva prioritat)."
                    ],
                    "nutrició": "No és mai la causa.",
                },
                "exploració": [
                    "Palpació L4-L5 dreta i L5-S1 dreta. Si en comparació amb l’esquerra fa més mal, si és la primera seria intestí prim, si és la segona sistema ginecològic femení/pròstata i per tant serà positiu.",
                    "Si les dues fan el mateix mal, palpar T9 a T12. Si la que li fa més mal és la T12, llavors és intestí prim, si és una de les altres, és sistema ginecològic femení/pròstata i per tant és positiu.",
                    "Si és pròstata/sistema ginecològic femení, llavors falta veure si és un estrès cronificat de vesícula biliar-intestí prim-ronyó o una emoció directament relacionada amb sistema ginecològic femení/pròstata. Per això li fem preguntes.",
                    "També podem buscar quins altres reflexes té a l’abdomen o a altres zones del cos, per exemple, dolor bessó, etc.",
                    "O si té altres símptomes de pròstata (com dolor en orinar) o d’intestí prim.", 
                    "Si és intestí prim li ha de fer mal la fosa ilíaca dreta.",
                    "Si continuem sense pistes, llavors fem més exploració (normalment no cal, només si tenim dubtes). Si és intestí prim la cama dreta estarà en rotació externa.", 
                    "O fer test d’allargament,  si allarga molt la cama dreta, llavors és intestí prim, cadena d’obertura.",
                    "O fer test d'escurçament i si la cama dreta escurça molt, llavors és sistema ginecològic femení/pròstata.", 
                    "Si és intestí prim, amb cames flexionades el genoll dret estarà més baix.", 
                    "I dret la cadera dreta més alta i assegut també.",
                    "Augment to cadena tancament dreta: psoes, adductors, vast intern..."
                    ],
                "a_tenir_en_compte": [
                    "El DIU pot donar problemes músculoesquelètics."
                    ]
            },
            
            "pròstata": {
                "símptomes": [
                    "Bloqueig L5 ERSd (extensió, inclinació i rotació dreta)",
                    "Hèrnia discal L5-S1 dreta",
                    "Dolor lumbar",
                    "Lumbàlgia de predomini dret",
                    "Sacroileitis dreta",
                    "Síndrome del piramidal dret (fals piramidal)",
                    "Trocanteritis dreta",
                    "Molèstia opressiva al perineu",
                    "Dolor a l'engonal dreta per tendinitis del psoes",
                    "Ciàtica dreta",
                    "Tendinitis de la pata de ganso (més freqüent dreta)",
                    "Fascitis plantar dreta",
                    "Tensió bessó dret",
                    "Dolor tibial posterior dret",
                    "Dolor bessó dret",
                    "Esperó calcani dret"
                    "dolor turmells",
                    "dolor peus",
                    "dolor lumbar",
                    "dolor genoll dret",
                    "dolor gluti dret"
                ],
                "altres_símptomes": [
                    "Micció costosa, entretallada, amb disminució del raig.",
                    "Molta dificultat per començar a orinar.",
                    "Necessitat d'aixecar-se a la nit a orinar diverses vegades.",
                    "Molèsties i irritació en orinar.",
                    "Orina residual o degoteig després d'orinar.",
                    "Sang a l'orina.",
                    "Ejaculació dolorosa.",
                    "Dolor suprapúbic, sensació d'opressió al perineu.",
                    "Urgència a la micció: necessitat d'orinar immediatament."
                ],
                "origen": {
                    "caràcter": "Persona introvertida, prudent, insegura, inadvertida, observadora, amb baixa autoestima i amb tendència a l'aïllament.",
                    "emocions": [
                        "Causa més comú: esgotament com a conseqüència de l’estrès, post estrès.",
                        "Autoestima danyada en l’àmbit de la parella, per sentir-se enganyat, abandonat, traït, o poc valorat, no sentir-se una prioritat per a la parella.",
                        "Si és per un fet puntual, això se soluciona normalment sent-ne conscient. Si és autoestima ja de caràcter, això necessita psicoteràpia.",
                        "Pors, inseguretat respecte als fills, patir molt per ells.",
                        "Tristor per malalties de llarga evolució amb deteriorament dels pares, habitual en gent que cuida dels pares.",
                        "Sentir-se degradat en el treball, o jubilació que no se la prenen bé perquè se senten inútils."
                    ],
                    "nutrició": "No és mai la causa."
                    
                                 },
                 "exploració": [
                    "Palpació L4-L5 dreta i L5-S1 dreta. Si en comparació amb l’esquerra fa més mal, si és la primera seria intestí prim, si és la segona sistema ginecològic femení/pròstata i per tant serà positiu.",
                    "Si les dues fan el mateix mal, palpar T9 a T12. Si la que li fa més mal és la T12, llavors és intestí prim, si és una de les altres, és sistema ginecològic femení/pròstata i per tant és positiu.",
                    "Si és pròstata/sistema ginecològic femení, llavors falta veure si és un estrès cronificat de vesícula biliar-intestí prim-ronyó o una emoció directament relacionada amb sistema ginecològic femení/pròstata. Per això li fem preguntes.",
                    "També podem buscar quins altres reflexes té a l’abdomen o a altres zones del cos, per exemple, dolor bessó, etc.",
                    "O si té altres símptomes de pròstata (com dolor en orinar) o d’intestí prim.", 
                    "Si és intestí prim li ha de fer mal la fosa ilíaca dreta.",
                    "Si continuem sense pistes, llavors fem més exploració (normalment no cal, només si tenim dubtes). Si és intestí prim la cama dreta estarà en rotació externa.", 
                    "O fer test d’allargament,  si allarga molt la cama dreta, llavors és intestí prim, cadena d’obertura.",
                    "O fer test d'escurçament i si la cama dreta escurça molt, llavors és sistema ginecològic femení/pròstata.", 
                    "Si és intestí prim, amb cames flexionades el genoll dret estarà més baix.", 
                    "I dret la cadera dreta més alta i assegut també.",
                    "Augment to cadena tancament dreta: psoes, adductors, vast intern..."
                    ],
                "a_tenir_en_compte": []                                
                },
                
            "estrès mantingut": {
                "símptomes": [],
                "altres_símptomes": [
                    "Acidosi metabòlica",
                    "Hipertensió arterial",
                    "Descalcificació",
                    "Obesitat",
                    "Esgotament"
                ],
                "origen": {
                    "caràcter": [],
                    "emocions": ["Estrès sostingut"],
                    "nutrició": []
                },
             "exploració":{
                "problemes_tren_superior": [
                    "Test en sedestació. Si li demanem extenxió i el cap s'inclina a la dreta i rota a l'esquerra, és que l'occipital està bloquejat en flexió. I si li demanem flexió i fa això mateix, llavors és que està bloquejat en extensió. (Atenció!! No és un test definitiu, ja que pot ser que la rotació i inclinació sigui el resultat d'una postura antiàlgica. Però en absència de dolor, en principi és bastant concloent).",
                    "T4 o T5 doloroses.",
                    "Espatlla dreta més baixa que l'esquerra.",
                    "Dolor a la palpació C1-C2 esquerra més que dreta."
            ],
                "problemes_tren_inferior": [
                    "En bipedestació i d'esquena, cresta ilíaca esquerra clarament més amunt que la dreta, giba lumbar dreta claríssima (aquest símptoma és inconstant), espatlla dreta més baixa.",
                    "En bipedestació, lleugera pèrdua de recurvàtum del genoll dret respecte a l'esquerra.",
                    "Dret i de cara, torsió pèlvica: EIAS dreta més baixa i EIAS esquerra més alta.",
                    "Assegut, cresta ilíaca esquerra igualment més alta que la dreta. En el test de flexió assegut de les EIPS, més mobilitat de la banda dreta (cal posar el dit just sota les EIPS per poder seguir bé el moviment).",
                    "En supí se segueix veient l'EIAS més alta que la dreta. Cama esquerra més curta que la dreta (falsa cama curta). (Per fer el test bé, abans apretar els genolls fins al tronc perquè no ens enganyi la postura). Augment del to rotatori de la cama esquerra, que per tant està més oberta que la dreta.",
                    "En supí, en el test de flexió de genolls (primer es porten els dos genolls cap a l'abdomen perquè no ens enganyi la postura i després es fa la comprovació amb les cuixes a 90 graus respecte al tronc i els genolls doblegats), genoll dret més alt que l'esquerre.",
                    "En prono: escoliosi lumbar i EIPS esquerra més alta.",
                    "En prono, angle lateral esquerra del sacre més cap al sostre, mentre que el dret està orientat a la camilla (sacre posterior esquerre). Normalment com a conseqüència el gluti esquerre puja una mica més que el dret."
            ]
}
               
            }

            }

    def normalitzar_simptomes(self, simptomes):
        """
        Converteix els símptomes introduïts pel pacient en termes mèdics equivalents,
        i afegeix els símptomes derivats de les zones corporals.
        """
       
        simptomes_normalitzats = set()
        for s in simptomes.split(","):
            s = s.strip().lower()
            
            # 🔹 Si el símptoma conté "mal ", afegim també la versió amb "dolor "
            if "mal " in s:
                simptomes_normalitzats.add(s.replace("mal ", "dolor "))

            # Afegim el símptoma original
            simptomes_normalitzats.add(s)

            # 🔹 Cas especial per "dret" i "esquerre"
            if "dret" in s or "esquerre" in s:
                base_sense_lateralitat = s.replace("dret", "").replace("esquerre", "").strip()

                # 🔹 També afegim la versió amb "dolor" si cal
                if "mal " in base_sense_lateralitat:
                    base_sense_lateralitat_dolor = base_sense_lateralitat.replace("mal ", "dolor ")
                    simptomes_normalitzats.add(base_sense_lateralitat_dolor)
                        
                if base_sense_lateralitat in self.zones_corpals:
                    simptomes_normalitzats.update(self.zones_corpals[s])

                elif base_sense_lateralitat in self.zones_corpals:
                    simptomes_normalitzats.update(self.zones_corpals[base_sense_lateralitat])
                    

            # Si és una zona corporal, afegim els símptomes associats
            if s in self.zones_corpals:
                simptomes_associats = self.zones_corpals[s]
                simptomes_normalitzats.update(simptomes_associats)
                

            # Si té un sinònim, afegim tant l’original com la versió mèdica
            if s in self.sinònims:
                simptomes_normalitzats.add(self.sinònims[s])

        simptomes_normalitzats = list(simptomes_normalitzats)
        return simptomes_normalitzats


    def identificar_organs_afectats(self, simptomes_usuari):
        """
        Identifica els òrgans afectats segons els símptomes introduïts.
        """
        simptomes_usuari = self.normalitzar_simptomes(simptomes_usuari)

        òrgans_afectats = []

        for òrgan, dades in self.vísceres.items():
            for símptoma in dades["símptomes"]:
                for part in simptomes_usuari:
                    # 🔹 Assegurem que comparem correctament
                    if part.strip().lower() in símptoma.strip().lower():
                        òrgans_afectats.append(òrgan)
                        break  # Si ja hi ha coincidència, passem al següent òrgan

        return list(set(òrgans_afectats))  # Eliminem duplicats
        
    # Obtenir fitoteràpia per organització i categoria
    def obtenir_fitoterapia_per_organ_i_categoria(self, tractaments_per_organ):
        """
        Classifica els tractaments de Fitoteràpia segons l'òrgan i la seva categoria,
        assegurant que cada òrgan es gestiona independentment.
        """
        tractaments_organitzats = {}

        for organ, tractaments in tractaments_per_organ.items():
            tractaments_organitzats[organ] = {}

            # Assignar tractaments a les categories dins del mateix òrgan
            for tractament in tractaments:
                categoria = self.classificacio_fitoterapia.get(tractament, "Sense categoria")

                # Afegim tractament a la categoria corresponent
                if categoria not in tractaments_organitzats[organ]:
                    tractaments_organitzats[organ][categoria] = []

                tractaments_organitzats[organ][categoria].append(tractament)

        return tractaments_organitzats


    def iniciar_qüestionari(self):
        global st
        st.title("👩‍⚕️ Assistent de diagnosi")

        # 🔹 Inicialitzar les variables de sessió si no existeixen
        if "respostes" not in st.session_state:
            st.session_state.respostes = {}

        if "òrgans_explorats" not in st.session_state:
            st.session_state["òrgans_explorats"] = []    

        if "preguntes_fetes" not in st.session_state:
            st.session_state.preguntes_fetes = set()
        
        if "exploracions_resultats" not in st.session_state:
            st.session_state["exploracions_resultats"] = {}

        # Inicialitzar respostes al session_state si no existeixen
        for pregunta in self.preguntes_inicials.keys():
            if pregunta not in st.session_state.respostes:
                st.session_state.respostes[pregunta] = ""

        # Preguntes inicials
        for pregunta, opcions in self.preguntes_inicials.items():
            if isinstance(opcions, list):
                index = (
                    opcions.index(st.session_state.respostes.get(pregunta, ""))
                    if pregunta in st.session_state.respostes and st.session_state.respostes[pregunta] in opcions
                    else None
                )
                resposta = st.radio(
                    pregunta,
                    opcions,
                    index=index,
                    key=f"pregunta_{pregunta}"
                )
            else:
                resposta = st.text_input(
                    pregunta,
                    value=st.session_state.respostes.get(pregunta, ""),  # Manté la resposta anterior
                    key=f"pregunta_{pregunta}"
                )

            # Guarda la resposta correctament
            st.session_state.respostes[pregunta] = resposta

            # 🔹 Preguntes dependents
            if resposta == "Sí":
                if pregunta == "Té lesions vertebrals diagnosticades?":
                    st.session_state.respostes["Detalls de les lesions vertebrals"] = st.text_input(
                        "Quines lesions vertebrals té?",
                        value=st.session_state.respostes.get("Detalls de les lesions vertebrals", ""),
                        key="detalls_lesions"
                    )
                elif pregunta == "Té malalties cròniques?":
                    st.session_state.respostes["Detalls de les malalties cròniques"] = st.text_input(
                        "Quines malalties cròniques té?",
                        value=st.session_state.respostes.get("Detalls de les malalties cròniques", ""),
                        key="detalls_malalties"
                    )
                    st.session_state.respostes["Medicaments que pren"] = st.text_input(
                        "Quins medicaments pren, si en pren algun?",
                        value=st.session_state.respostes.get("Medicaments que pren", ""),
                        key="medicaments"
                    )
                elif pregunta == "Té al·lèrgies?":
                    st.session_state.respostes["Detalls de les al·lèrgies"] = st.text_input(
                        "Quines al·lèrgies té?",
                        value=st.session_state.respostes.get("Detalls de les al·lèrgies", ""),
                        key="detalls_al·lèrgies"
                    )

        # Definir òrgans afectats
        òrgans_afectats = []
        if "Quins símptomes músculoesquelètics té?" in st.session_state.respostes:
            simptomes_usuari = st.session_state.respostes["Quins símptomes músculoesquelètics té?"].strip().lower()
            if "simptomes_manuals" not in st.session_state:
                st.session_state["simptomes_manuals"] = st.session_state.respostes["Quins símptomes músculoesquelètics té?"]

            # Guarda només el que ha escrit l'usuari
            st.session_state.respostes["Quins símptomes músculoesquelètics té?"] = st.session_state["simptomes_manuals"]

            # 🔹 Separar els símptomes introduïts per l'usuari
            simptomes_usuari = simptomes_usuari.split(",")

            # 🔹 Identificar si ha escrit una zona corporal en comptes d'un símptoma específic
            simptomes_sugerits = []
            for s in simptomes_usuari:
                s = s.strip()
                if s in self.zones_corpals:
                    simptomes_sugerits.extend(self.zones_corpals[s])  # Afegim els símptomes associats

            # 🔹 Inicialitzar `simptomes_confirmats` per evitar l'error
            simptomes_confirmats = []

            # 🔹 Si detectem que ha escrit una zona corporal, oferim selecció guiada
            if simptomes_sugerits:

                # 🔹 Afegir els símptomes seleccionats a la llista final
                if simptomes_confirmats:
                    simptomes_usuari.extend(simptomes_confirmats)

            # 🔹 Aplicar normalització de símptomes (sinònims + zones corporals)
            simptomes_sugerits = self.normalitzar_simptomes(st.session_state["simptomes_manuals"])

            # 🔹 Guardem els símptomes processats a session_state
            st.session_state.respostes["Quins símptomes músculoesquelètics té?"] = ", ".join(simptomes_usuari)
            
            # 🔹 Assegurem que `identificar_organs_afectats()` rebi els símptomes correctes
            if simptomes_usuari and any(s.strip() for s in simptomes_usuari):  # Comprova que no està buit
                òrgans_afectats = self.identificar_organs_afectats(", ".join(simptomes_usuari))
            else:
                òrgans_afectats = []  # No hi ha òrgans afectats fins que s'introdueixin símptomes

            for s in simptomes_usuari:
                s = s.strip()

                # 🔹 Afegir els símptomes seleccionats a la llista final
                if simptomes_confirmats:
                    simptomes_usuari.extend(simptomes_confirmats)

            # 🔹 Guardem els símptomes definitius a session_state per continuar el diagnòstic
            simptomes_usuari = ", ".join(simptomes_usuari)  # 🛠️ Convertim la llista a una cadena

            st.session_state.respostes["Quins símptomes músculoesquelètics té?"] = simptomes_usuari

            # 🔹 Assegurem que estem passant una cadena a la funció de detecció d'òrgans
            if simptomes_usuari:
                òrgans_afectats = self.identificar_organs_afectats(simptomes_usuari)
                if not òrgans_afectats:
                    st.warning(
                        "⚠️ No soc capaç de trobar correspondències amb cap òrgan a partir dels símptomes introduïts."
                    )

                    st.markdown(
                        """
                        🔹 **Suggeriments per millorar la introducció de símptomes:**
                        - Revisa que no hagis comès alguna falta ortogràfica o de picat.
                        - Escriu **cada símptoma separat per comes** (exemple: `"dolor cervical, mareig, mal de cap"`).
                        - **Evita articles i preposicions** (❌ `"el meu genoll fa mal"` ➝ ✅ `"dolor genoll"`).
                        - Usa **termes senzills** i evita frases llargues.
                        - Si el símptoma afecta una zona específica, indica si és **dret o esquerre** (exemple: `"dolor espatlla dreta"`).
                        """
                    )

        # Afegim "Estrès mantingut" només si no hi és
        if òrgans_afectats and "estrès mantingut" not in òrgans_afectats:
            òrgans_afectats.append("estrès mantingut")

        # Excloure òrgans segons el sexe seleccionat
            sexe_usuari = st.session_state.respostes.get("És dona o home?", "")

            if sexe_usuari == "Home":
               òrgans_afectats = [òrgan for òrgan in òrgans_afectats if òrgan != "sistema ginecològic femení"]
            elif sexe_usuari == "Dona":
                    òrgans_afectats = [òrgan for òrgan in òrgans_afectats if òrgan != "pròstata"]

        # Mostrar els òrgans potencialment implicats
        if òrgans_afectats:
            st.subheader("📌 Òrgans potencialment implicats o altres causes")
            
            # ✅ Pregunta si es vol veure la llista de símptomes músculoesquelètics
            mostrar_símptomes = st.checkbox("Vols veure la llista completa de símptomes músculoesquelètics associats a aquests òrgans?")

            if mostrar_símptomes:
                st.subheader("📋 Símptomes músculoesquelètics per òrgan")
                for òrgan in òrgans_afectats:
                    st.write(f"### {òrgan.capitalize()}")
                    simptomes_musculoesqueletics = self.vísceres.get(òrgan, {}).get("símptomes", [])
                    if simptomes_musculoesqueletics:
                        for s in simptomes_musculoesqueletics:
                            st.write(f"🔹 {s}")
                    else:
                        st.write("No hi ha símptomes músculoesquelètics definits per aquest òrgan.")

            # Identificar tractaments de fitoteràpia associats als òrgans detectats
            tractaments_per_organ = {}

            for organ in òrgans_afectats:
                if organ in self.vísceres and "fitoteràpia" in self.vísceres[organ]:
                    tractaments_per_organ[organ] = self.vísceres[organ]["fitoteràpia"]

            # Classificar els tractaments segons l'òrgan i la seva categoria
            fitoterapia_classificada = self.obtenir_fitoterapia_per_organ_i_categoria(tractaments_per_organ)

            # 🔹 Mostrar els tractaments classificats per òrgan i categoria
            for organ, categories in fitoterapia_classificada.items():
                st.markdown(f"### 🏷 {organ.capitalize()}")  # Mostra l'òrgan com a títol

                for categoria, tractaments in categories.items():
                    if categoria == "Sense categoria":
                        st.subheader("➡ Tractaments sense classificació")
                    else:
                        st.subheader(f"➡ {categoria}")  # Mostra la categoria dins de l'òrgan

                    for tractament in tractaments:
                        st.write(f"🔹 {tractament}")

            # ✅ Afegim aquí l'associació de símptomes amb patologies
            if simptomes_usuari:
                if isinstance(simptomes_usuari, str):  # Si és una cadena, la convertim en una llista
                    simptomes_usuari = simptomes_usuari.split(", ")

                # Trobar el símptoma que inclou totes les possibilitats
                simptoma_més_complet = None
                for s in simptomes_usuari:
                    s = s.strip()  # Elimina espais innecessaris
                    if s:  # Evita processar elements buits
                        simptomes_associats = self.normalitzar_simptomes(s)
                        # Triem el primer símptoma que té més d'una associació com el més complet
                        if len(simptomes_associats) > 1:
                            simptoma_més_complet = (s, simptomes_associats)
                            break  # Ens aturem a la primera coincidència

                # Mostrem només el símptoma amb més associacions
                if simptoma_més_complet:
                    st.write(f"✅ Associant {simptoma_més_complet[0]} amb {simptoma_més_complet[1]}")

            for òrgan in òrgans_afectats:
                st.write(f"➡ **{òrgan.capitalize()}**")
        else:
            st.subheader("⚠️ No s'han trobat coincidències amb cap òrgan")
            st.write("Si us plau, revisa els símptomes introduïts i assegura't que estan ben escrits.")

       
        # Formulari per seleccionar altres símptomes
        st.markdown("---")  # 🔹 SEPARACIÓ VISUAL
        st.subheader("📋 Comprovació d'altres símptomes")
        símptomes_adicionals = {}

        # Afegir els altres símptomes de cada òrgan detectat
        for òrgan in òrgans_afectats:
            if òrgan in self.vísceres and "altres_símptomes" in self.vísceres[òrgan]:
                st.write(f"**{òrgan.capitalize()}**")
                símptomes_adicionals[òrgan] = []
                for símptoma in self.vísceres[òrgan]["altres_símptomes"]:
                    if st.checkbox(símptoma, key=f"{òrgan}_{símptoma}"):
                        símptomes_adicionals[òrgan].append(símptoma)

        # Guardar els altres símptomes marcats
        st.session_state["símptomes_adicionals"] = símptomes_adicionals
        
        # 🔹 Inicialitzar la clau "símptomes_estrès" si no existeix
        if "símptomes_estrès" not in st.session_state:
            st.session_state["símptomes_estrès"] = []


        # Afegir estil personalitzat per al botó de confirmació
        st.markdown("""
            <style>
                div.stButton > button:first-child {
                    background-color: #28a745;
                    color: white;
                    border-radius: 5px;
                    padding: 10px 20px;
                    font-size: 16px;
                }
                div.stButton > button:first-child:hover {
                    background-color: #218838;
                    color: white;
                }
            </style>
        """, unsafe_allow_html=True)

        # Botó de confirmació per als altres símptomes
        if st.button("Confirmar", key="confirmar_símptomes"):
            òrgans_seleccionats = [
                òrgan for òrgan, símptomes_marcats in símptomes_adicionals.items() if símptomes_marcats
            ]
            
            # Els òrgans que no s'han seleccionat però ja estaven detectats abans
            òrgans_pendents = [òrgan for òrgan in òrgans_afectats if òrgan not in òrgans_seleccionats]

            st.subheader("📋 Resultats de la selecció d'altres símptomes")

            if òrgans_seleccionats:
                st.write("✅ Els altres símptomes suggereixen que els òrgans implicats són:")
                for òrgan in òrgans_seleccionats:
                    st.write(f"- **{òrgan.capitalize()}**")
            else:
                st.write("⚠️ No s'ha seleccionat cap altre símptoma.")

            if òrgans_pendents:
                st.write("ℹ Igualment, no podem descartar aquests òrgans:")
                for òrgan in òrgans_pendents:
                    st.write(f"- **{òrgan.capitalize()}**")

            # Si no s'ha seleccionat cap altre símptoma, mantenim els òrgans del pas anterior
            if not òrgans_seleccionats:
                òrgans_seleccionats = òrgans_afectats  

            # Assegurem que en el següent pas es mantinguin els òrgans detectats inicialment
            st.session_state["òrgans_definitius"] = list(set(òrgans_afectats + òrgans_seleccionats))

        # **Només mostrar la següent secció si ja s'han confirmat els altres símptomes**
        if "òrgans_definitius" in st.session_state and st.session_state["òrgans_definitius"]:
            st.markdown("---")  # 🔹 SEPARACIÓ VISUAL
            st.subheader("🧩 Potencials origens de la disfunció")

            origens_adicionals = {}

            for òrgan in st.session_state["òrgans_definitius"]:
                if "origen" in self.vísceres[òrgan]:
                    st.write(f"**{òrgan.capitalize()}**")
                    if òrgan == "estrès mantingut":
                        st.write("🔵 **Origen:** Estrès mantingut")
                        if st.checkbox("Estrès mantingut", key=f"{òrgan}_origen"):
                            origens_adicionals.setdefault(òrgan, {"origen": []})["origen"].append("Estrès mantingut")
                        continue  # ❗ Saltem la resta del codi per evitar errors
                    origens_adicionals[òrgan] = {
                        "caràcter": [],
                        "emocions": [],
                        "nutrició": []
                    }

                    # Selecció de caràcter
                    if "caràcter" in self.vísceres[òrgan]["origen"]:
                        st.write("🟢 **Caràcter**")
                        if st.checkbox(self.vísceres[òrgan]["origen"]["caràcter"], key=f"{òrgan}_caràcter"):
                            origens_adicionals[òrgan]["caràcter"].append(self.vísceres[òrgan]["origen"]["caràcter"])

                    # Selecció d'emocions
                    if "emocions" in self.vísceres[òrgan]["origen"]:
                        st.write("🔵 **Emocions**")
                        for emocio in self.vísceres[òrgan]["origen"]["emocions"]:
                            if st.checkbox(emocio, key=f"{òrgan}_emocions_{emocio}"):
                                origens_adicionals[òrgan]["emocions"].append(emocio)

                    # Selecció de nutrició
                    if "nutrició" in self.vísceres[òrgan]["origen"]:
                        st.write("🟠 **Nutrició**")
                        if st.checkbox(self.vísceres[òrgan]["origen"]["nutrició"], key=f"{òrgan}_nutrició"):
                            origens_adicionals[òrgan]["nutrició"].append(self.vísceres[òrgan]["origen"]["nutrició"])

            # Inicialitzar les claus de session_state si no existeixen
            if "no_encaixa_origens" not in st.session_state:
                st.session_state["no_encaixa_origens"] = False

            if "observacions_origens" not in st.session_state:
                st.session_state["observacions_origens"] = ""
                
            # Afegir espai abans de l'opció "No encaixa cap d'aquests potencials origens"
            st.markdown('<div style="margin-top: 10px; margin-bottom: 10px; border-top: 1px dashed lightgray;"></div>', unsafe_allow_html=True)

            # Opció "No encaixa cap d'aquests potencials origens"
            no_encaixa = st.checkbox("❌ No encaixa cap d'aquests potencials origens", 
                                     value=st.session_state["no_encaixa_origens"], 
                                     key="no_encaixa_origens")

            # Casella d'observacions
            observacions = st.text_area("📝 Observacions", 
                                        value=st.session_state["observacions_origens"], 
                                        key="observacions_origens")

            # Guardar els potencials origens seleccionats
            st.session_state["origens_adicionals"] = origens_adicionals


            # ✅ Botó de confirmació per als potencials origens
            if st.button("Confirmar", key="confirmar_origens"):
                st.subheader("🔍 Resultats seleccionats")

                if any(st.session_state["origens_adicionals"].values()):
                    st.write("📌 **Origens seleccionats:**")
                    for òrgan, dades in st.session_state["origens_adicionals"].items():
                        if any(dades.values()):  # Comprovar si hi ha algun origen seleccionat
                            st.write(f"➡ **{òrgan.capitalize()}**")
                            if "caràcter" in dades and dades["caràcter"]:
                                st.write(f"🟢 **Caràcter:** {', '.join(dades['caràcter'])}")

                            if "emocions" in dades and dades["emocions"]:
                                st.write(f"🔵 **Emocions:** {', '.join(dades['emocions'])}")

                            if "nutrició" in dades and dades["nutrició"]:
                                st.write(f"🟠 **Nutrició:** {', '.join(dades['nutrició'])}")


                if st.session_state.get("no_encaixa_origens", False):
                    st.write("❌ **S'ha marcat 'No encaixa cap d'aquests potencials origens'.**")

                if st.session_state.get("observacions_origens", ""):
                    st.write(f"📝 **Observacions:** {st.session_state['observacions_origens']}")
 
            # 🔎 **Exploració clínica**
            st.markdown("---")  # 🔹 SEPARACIÓ VISUAL
            if "òrgans_definitius" in st.session_state:
                st.subheader("🔎 Exploració clínica")

                # ✅ Opció multiresposta per seleccionar exploracions a realitzar
                exploracions_disponibles = [
                    "Fetge", "Vesícula biliar", "Estómac", "Sistema circulatori/cor",
                    "Pàncrees", "Intestí prim", "Colon irritable", "Ronyó", "Bufeta",
                    "Sistema ginecològic femení", "Pròstata",
                    "Estrès mantingut tren superior", "Estrès mantingut tren inferior",
                    "No fem exploració"
                ]

                # Inicialitzar variables si no existeixen
                if "seleccio_exploracions" not in st.session_state:
                    st.session_state["seleccio_exploracions"] = []

                if "exploracions_confirmades" not in st.session_state:
                    st.session_state["exploracions_confirmades"] = []

                if "exploracions_temporals" not in st.session_state:
                    st.session_state["exploracions_temporals"] = {}

                # Multiselect per escollir exploracions
                exploracions_seleccionades = st.multiselect(
                    "Selecciona les exploracions pertinents:",
                    exploracions_disponibles,
                    default=st.session_state["seleccio_exploracions"],  
                    key="seleccio_exploracions"
                )

                # ✅ Guardar selecció automàticament
                st.session_state["exploracions_confirmades"] = exploracions_seleccionades

                # 🔎 **Exploració confirmada**
                if st.session_state["exploracions_confirmades"]:
                    if "No fem exploració" in st.session_state["exploracions_confirmades"]:
                        st.warning("ℹ️ No s'ha seleccionat cap exploració a realitzar.")
                    else:
                        st.subheader("")

                        # Diccionari per normalitzar noms d'òrgans
                        correspondencia = {
                            "Sistema circulatori/cor": "sistema circulatori i cor",
                            "Colon irritable": "colon irritable",
                            "Intestí prim": "intestí prim",
                            "Vesícula biliar": "vesícula biliar",
                            "Sistema ginecològic femení": "sistema ginecològic femení",
                            "Estrès mantingut tren superior": "estrès mantingut",
                            "Estrès mantingut tren inferior": "estrès mantingut"
                        }

                        for òrgan in st.session_state["exploracions_confirmades"]:
                            òrgan_clau = correspondencia.get(òrgan, òrgan.lower())

                            st.write(f"### {òrgan.capitalize()}")

                            exploracio = self.vísceres.get(òrgan_clau, {}).get("exploració", None)

                            # 📌 Cas especial per estrès mantingut
                            if òrgan in ["Estrès mantingut tren superior", "Estrès mantingut tren inferior"]:
                                subcategoria = "problemes_tren_superior" if "superior" in òrgan else "problemes_tren_inferior"
                                exploracio = self.vísceres.get("estrès mantingut", {}).get("exploració", {}).get(subcategoria, None)

                            # 🚨 **Si no hi ha exploració, mostrar avís**
                            if not exploracio:
                                st.warning(f"⚠️ Aquest òrgan no té exploració associada: **{òrgan}**")
                                continue  # Passar al següent òrgan

                            # 🔍 **Mostra els punts d'exploració**
                            if isinstance(exploracio, list):  
                                for punt in exploracio:
                                    st.markdown(f"**{punt}**")  
                                    col1, col2, col3 = st.columns(3)  
                                    with col1:
                                        pos = st.checkbox("Positiu", key=f"{òrgan_clau}_{punt}_pos")
                                    with col2:
                                        neg = st.checkbox("Negatiu", key=f"{òrgan_clau}_{punt}_neg")
                                    with col3:
                                        inconcl = st.checkbox("Sense resultats concloents", key=f"{òrgan_clau}_{punt}_inconcl")

                                    # Guardar el resultat seleccionat automàticament
                                    if pos:
                                        st.session_state["exploracions_temporals"][f"{òrgan_clau}_{punt}"] = "Positiu"
                                    elif neg:
                                        st.session_state["exploracions_temporals"][f"{òrgan_clau}_{punt}"] = "Negatiu"
                                    elif inconcl:
                                        st.session_state["exploracions_temporals"][f"{òrgan_clau}_{punt}"] = "Sense resultats concloents"

                                # Actualitzar exploracions_resultats directament
                                st.session_state["exploracions_resultats"] = st.session_state["exploracions_temporals"]

                        # 📌 **Afegir un apartat per posar observacions**
                        st.session_state["observacions_resultats"] = st.text_area(
                            "📝 Observacions addicionals",
                            value=st.session_state.get("observacions_resultats", ""),
                            key="observacions_exploracio")

                        # 📋 **Mostrar resultats finals immediatament**
                        st.subheader("📋 Resultats finals de l'exploració")
                        if st.session_state["exploracions_resultats"]:
                            for exploracio, resultat in st.session_state["exploracions_resultats"].items():
                                st.write(f"🔹 **{exploracio.replace('_', ' ')}**: {resultat}")
                        else:
                            st.write("⚠️ No s'ha marcat cap resultat per a l'exploració.")

                        # Mostrar les observacions si s'han escrit
                        if st.session_state["observacions_resultats"]:
                            st.subheader("📝 Observacions registrades")
                            st.write(st.session_state["observacions_resultats"])


            # 📌 **A TENIR EN COMPTE**
            if "exploracions_confirmades" in st.session_state and st.session_state["exploracions_confirmades"]:
                st.subheader("📌 A tenir en compte")

                cap_informacio = True  # Variable per verificar si hi ha informació a mostrar

                for òrgan_clau in st.session_state["exploracions_confirmades"]:
                    # 🔍 **Normalitzem els noms perquè coincideixin amb la llibreria**
                    òrgan_clau = òrgan_clau.lower()

                    if òrgan_clau in self.vísceres:  # ✅ Comprovem si l'òrgan està a la llibreria
                        dades_organ = self.vísceres[òrgan_clau]
                        a_tenir_en_compte = dades_organ.get("a_tenir_en_compte", [])

                        if a_tenir_en_compte:  # ✅ Només mostrem si hi ha informació
                            st.write(f"### {òrgan_clau.capitalize()}")
                            for punt in a_tenir_en_compte:
                                st.write(f"- {punt}")
                            cap_informacio = False  # Hem trobat informació, canviem el valor de la variable

                # ❗ Si cap òrgan té informació rellevant, mostrem el missatge per defecte
                if cap_informacio:
                    st.write("ℹ No hi ha informació addicional a tenir en compte per als òrgans seleccionats.")
            
            import streamlit as st
            import docx 
            from docx import Document
            from io import BytesIO

            def llegir_tractaments(nom_document):
                """
                Llegeix el document Word i extreu els tractaments correctament associats als òrgans.
                Detecta els òrgans en MAJÚSCULES i els assigna als tractaments corresponents.
                Manté les seccions ("A la consulta", "A casa", etc.), però dins de cada òrgan.
                """
                from docx import Document
                import re

                doc = Document(nom_document)
                tractaments = {}
                organ_actual = None  # Per saber a quin òrgan pertany cada secció
                seccio_actual = None
                tractament_actual = None

                for par in doc.paragraphs:
                    text = par.text.strip()

                    if not text:
                        continue  # Ometre línies buides

                    # 🔹 **Detectar ÒRGANS (en MAJÚSCULES)**
                    if text.isupper() and len(text) > 3:  # Evita que es confongui amb seccions curtes
                        organ_actual = text
                        tractaments.setdefault(organ_actual, {})  # Crear entrada per a l'òrgan
                        seccio_actual = None  # Reiniciar la secció
                        tractament_actual = None  # Reiniciar el tractament
                        continue

                    # 🔹 **Detectar seccions com "A la consulta", "A casa", etc.**
                    if text.startswith("###"):
                        if not organ_actual:
                            continue  # Evita errors si no hi ha un òrgan definit

                        seccio_actual = text.replace("###", "").strip()
                        tractaments[organ_actual].setdefault(seccio_actual, {})  # Crear secció dins l'òrgan
                        tractament_actual = None  # Reiniciar tractament
                        continue

                    # 🔸 **Detectar els tractaments correctes (els que tenen >>> i <<<)**
                    match = re.match(r">>>(.*?)<<<", text)
                    if match:
                        if not organ_actual or not seccio_actual:
                            continue  # Evita errors si no hi ha una estructura correcta

                        tractament_actual = match.group(1).strip()
                        tractaments[organ_actual][seccio_actual].setdefault(tractament_actual, "")
                        continue

                    # ✅ **Afegir el contingut al tractament corresponent**
                    if organ_actual and seccio_actual and tractament_actual:
                        tractaments[organ_actual][seccio_actual][tractament_actual] += text + "\n"

                return tractaments
             
            # 📥 Carregar tractaments des del document Word
            tractaments = llegir_tractaments("documents/Tractaments.docx")

            # Crear diccionari per organitzar els tractaments
            tractaments_organitzats = {}

            # Obtenir i organitzar tractaments per òrgan i categoria
            for organ, seccions in tractaments.items():
                tractaments_organitzats[organ] = {}  # Inicialitzem per cada òrgan

                # Primer, mostrem les seccions comunes
                for seccio, tractaments_seccio in seccions.items():
                    if seccio in ["A casa", "A la consulta", "Nutrició", "Fitoteràpia"]:
                        tractaments_organitzats[organ][seccio] = tractaments_seccio
                        continue  # Evitem processar altres seccions més endavant

                # Ara processem Fitoteràpia categories com l'última secció
                if "Fitoteràpia categories" in seccions:
                    if "Fitoteràpia categories" not in tractaments_organitzats[organ]:
                        tractaments_organitzats[organ]["Fitoteràpia categories"] = {}

                    # Classifiquem els tractaments dins de Fitoteràpia categories
                    for tractament in seccions["Fitoteràpia categories"]:
                        tractament_net = tractament.strip().replace("\u200b", "").replace("\xa0", " ")

                        # Busquem la categoria corresponent dins del diccionari de classificació
                        categoria = self.classificacio_fitoterapia.get(tractament_net)

                        if not categoria:
                            categoria = "Sense categoria"

                        if categoria not in tractaments_organitzats[organ]["Fitoteràpia categories"]:
                            tractaments_organitzats[organ]["Fitoteràpia categories"][categoria] = []

                        tractaments_organitzats[organ]["Fitoteràpia categories"][categoria].append(tractament_net)


            # 🔹 Selecció manual dels òrgans per obtenir tractaments
            st.markdown("---")  # 🔹 SEPARACIÓ VISUAL
            st.subheader("📌 Selecciona els òrgans per als quals vols veure el tractament:")
            òrgans_disponibles = list(tractaments.keys())  # Òrgans disponibles al document
            òrgans_seleccionats = st.multiselect("Tria els òrgans:", òrgans_disponibles, key="tractaments_seleccionats")

            # 🔎 **Mostrar tractament per als òrgans explorats**
            st.subheader("💊 Tractament recomanat")

            if "tractaments_seleccionats" in st.session_state and st.session_state["tractaments_seleccionats"]:
                tractaments_a_mostrar = {organ: tractaments_organitzats.get(organ, {}) for organ in st.session_state["tractaments_seleccionats"]}

                for organ, seccions in tractaments_a_mostrar.items():
                    st.subheader(f"🦠 Tractament per a {organ}")

                    for seccio, tractaments_seccio in seccions.items():
                        if not tractaments_seccio:  # Evitar mostrar seccions buides
                            continue  

                        if seccio in ["A casa", "A la consulta", "Nutrició", "Fitoteràpia"]:
                            if seccio == "A la consulta":
                                st.subheader("💼 A la consulta")
                            elif seccio == "Nutrició":
                                st.subheader("🥗 Nutrició")
                            elif seccio == "Fitoteràpia":
                                st.subheader("🌱 Fitoteràpia")
                            else:
                                st.subheader(f"🏡 {seccio}")  # Per la resta de seccions

                            for i, (tractament, contingut) in enumerate(tractaments_seccio.items()):  # Afegim un índex únic
                                with st.expander(f"▶ {tractament}"):
                                    seleccionat = st.checkbox(f"✅ Seleccionar aquest tractament", key=f"{organ}_{seccio}_{tractament}_{i}")  # Clau única

                                    if isinstance(contingut, str):  # Només dividir si és una cadena
                                        paragrafs = contingut.split("\n")
                                    else:
                                        paragrafs = [contingut]  # Si és un diccionari, el mantenim tal qual

                                    for par in paragrafs:
                                        if isinstance(par, str) and par.strip():  # Només aplicar .strip() si és una cadena
                                            st.markdown(par.strip(), unsafe_allow_html=True)
                                            
                                    # 🔹 Importar `os` dins del bloc on s'utilitza
                                    import os  
                                    
                                    # 🔹 Buscar la imatge associada només per "A la consulta" i "A casa"
                                    if seccio in ["A la consulta", "A casa"]:
                                        image_filename_png = f"{tractament.replace(' ', '_')}.png"
                                        image_filename_jpg = f"{tractament.replace(' ', '_')}.jpg"

                                        image_path_png = os.path.join("imatges", image_filename_png)
                                        image_path_jpg = os.path.join("imatges", image_filename_jpg)

                                        # ✅ Si existeix una imatge, la mostrem
                                        if os.path.exists(image_path_png):
                                            st.image(image_path_png, caption=f"Imatge per {tractament}")
                                        elif os.path.exists(image_path_jpg):
                                            st.image(image_path_jpg, caption=f"Imatge per {tractament}")

                                    # 🔹 Gestionar tractaments seleccionats
                                    if "tractaments_seleccionats_usuari" not in st.session_state:
                                        st.session_state["tractaments_seleccionats_usuari"] = []

                                    if seleccionat:
                                        # Afegir només si no està ja afegit
                                        if (organ, seccio, tractament) not in st.session_state["tractaments_seleccionats_usuari"]:
                                            st.session_state["tractaments_seleccionats_usuari"].append((organ, seccio, tractament))
                                    else:
                                        # **Eliminar si el checkbox es desmarca**
                                        if (organ, seccio, tractament) in st.session_state["tractaments_seleccionats_usuari"]:
                                            st.session_state["tractaments_seleccionats_usuari"].remove((organ, seccio, tractament))

                        # 🔹 Mostrem Fitoteràpia categories sense títol de secció
                        elif seccio == "Fitoteràpia categories":
                            for categoria, tractaments_categoria in tractaments_seccio.items():
                                with st.expander(f"📌 {categoria.capitalize()}"):
                                    # Mostrem només els noms dels tractaments dins de cada categoria
                                    for i, tractament in enumerate(tractaments_categoria):
                                        seleccionat = st.checkbox(f"✅ {tractament}", key=f"{organ}_{seccio}_{categoria}_{tractament}_{i}")

                                        # ✅ Inicialitzar session_state si no existeix
                                        if "tractaments_seleccionats_usuari" not in st.session_state:
                                            st.session_state["tractaments_seleccionats_usuari"] = []

                                        if seleccionat:
                                            # **Evitar duplicats abans d'afegir**
                                            if (organ, seccio, tractament) not in st.session_state["tractaments_seleccionats_usuari"]:
                                                st.session_state["tractaments_seleccionats_usuari"].append((organ, seccio, tractament))
                                        else:
                                            # **Eliminar només si ja existia**
                                            if (organ, seccio, tractament) in st.session_state["tractaments_seleccionats_usuari"]:
                                                st.session_state["tractaments_seleccionats_usuari"].remove((organ, seccio, tractament))

            else:
                st.write("❗ Selecciona un òrgan per veure el tractament.")
                
            # 📝 Caselles per a observacions addicionals
            altres_tractaments_consulta = st.text_area("🛠️ Altres tractaments aplicats a la consulta", key="altres_tractaments_consulta")
            altres_recomanacions_casa = st.text_area("🏠 Altres recomanacions a casa", key="altres_recomanacions_casa")

            def generar_documents_word(nom_pacient, tractaments_a_mostrar):
                """
                Genera dos documents Word:
                1. Document per al pacient amb el tractament complet.
                2. Document per a la consulta amb un historial resumit.
                """

                document_pacient = Document()
                document_consulta = Document()

                document_pacient.add_heading(f"Pla de tractament per a {nom_pacient}", level=1)
                document_consulta.add_heading(f"Historial de consulta - {nom_pacient}", level=1)

                for organ, tractament in tractaments_a_mostrar.items():
                    document_pacient.add_heading(organ, level=2)
                    document_consulta.add_heading(organ, level=2)

                    if isinstance(tractament, dict) and tractament:  # Si hi ha tractament
                        for seccio, contingut in tractament.items():
                            document_pacient.add_heading(seccio, level=3)
                            document_pacient.add_paragraph(contingut)

                            document_consulta.add_heading(seccio, level=3)
                            document_consulta.add_paragraph(contingut)
                    else:
                        document_pacient.add_paragraph("ℹ No hi ha tractament disponible per a aquest òrgan.")
                        document_consulta.add_paragraph("ℹ No hi ha informació disponible.")

                pacient_doc_path = f"Pla_tractament_{nom_pacient}.docx"
                consulta_doc_path = f"Historial_consulta_{nom_pacient}.docx"

                document_pacient.save(pacient_doc_path)
                document_consulta.save(consulta_doc_path)

                return pacient_doc_path, consulta_doc_path
                
            print("🔍 DEBUG - Dades de session_state abans de generar el document:")
            for clau, valor in st.session_state.items():
                print(f"{clau}: {valor}")
                
            st.markdown("---")  # 🔹 SEPARACIÓ VISUAL abans de la generació de documents    

            def generar_historial(pacient, dades_pacient, tractaments_seleccionats):
                """
                Genera un document Word amb l'historial complet del pacient.
                """

                doc = Document()
                from docx.shared import Pt

                style = doc.styles['Normal']
                style.font.name = 'Arial'
                style.font.size = Pt(11)

                # Assegurar que Word reconeix la font (a vegades cal aquesta línia addicional)
                for para in doc.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)

                doc.add_heading(f"Historial de {pacient}", level=1)

                respostes = dades_pacient.get("respostes", {})  # 🔹 Agafar totes les respostes correctament
                
                òrgans_definitius = st.session_state.get("òrgans_definitius", [])
                if òrgans_definitius:
                    origens_emocionals = origens_adicionals.get(òrgans_definitius[0], {}).get("emocions", [])
                else:
                    origens_emocionals = []


                # 📌 1️⃣ Nom i gènere
                doc.add_paragraph(f"👤 **Nom del pacient:** {respostes.get('Nom del pacient', 'No especificat')}")
                doc.add_paragraph(f"🧑‍⚕️ **Gènere:** {respostes.get('És dona o home?', 'No especificat')}")

                # 📌 2️⃣ Símptomes musculoesquelètics
                doc.add_heading("Símptomes musculoesquelètics", level=2)
                doc.add_paragraph(respostes.get("Quins símptomes músculoesquelètics té?", "No especificat"))
                
                # 📌 2️⃣ Explicació dels símptomes
                doc.add_heading("Explicació dels símptomes", level=2)
                doc.add_paragraph(respostes.get("Explicació dels símptomes", "No especificat"))

                # 📌 3️⃣ Des de quan?
                doc.add_heading("Des de quan?", level=2)
                doc.add_paragraph(respostes.get("Des de quan?", "No especificat"))

                # 📌 4️⃣ En quins moments del dia li passa?
                doc.add_heading("En quins moments del dia li passa?", level=2)
                doc.add_paragraph(respostes.get("En quins moments del dia li passa? Quan o on es troba quan li passa en general?", "No especificat"))

                # 📌 5️⃣ Lesions vertebrals
                doc.add_heading("Lesions vertebrals", level=2)
                doc.add_paragraph(respostes.get("Té lesions vertebrals diagnosticades?", "No especificat"))
                doc.add_paragraph(respostes.get("Detalls de les lesions vertebrals", "No especificat"))

                # 📌 6️⃣ Malalties cròniques
                doc.add_heading("Malalties cròniques", level=2)
                doc.add_paragraph(respostes.get("Té malalties cròniques?", "No especificat"))
                doc.add_paragraph(respostes.get("Detalls de les malalties cròniques", "No especificat"))
                doc.add_paragraph(f"💊 Medicaments que pren: {respostes.get('Medicaments que pren', 'No especificat')}")

                # 📌 7️⃣ Al·lèrgies
                doc.add_heading("Al·lèrgies", level=2)
                doc.add_paragraph(respostes.get("Té al·lèrgies?", "No especificat"))
                doc.add_paragraph(respostes.get("Detalls de les al·lèrgies", "No especificat"))
                
                # Obtenir òrgans_definitius de manera segura
                organs_definitius = st.session_state.get("òrgans_definitius", [])

                # 📌 Altres símptomes
                doc.add_heading("Altres símptomes", level=2)
                altres_símptomes = símptomes_adicionals.get(organs_definitius[0], []) if organs_definitius else []
                if altres_símptomes:
                    for símptoma in altres_símptomes:
                        doc.add_paragraph(símptoma)
                else:
                    doc.add_paragraph("No especificat")              

                # 📌 Origens emocionals
                doc.add_heading("Origens emocionals", level=2)
                origens_emocionals = origens_adicionals.get(òrgans_definitius[0], {}).get("emocions", []) if òrgans_definitius else []
                if origens_emocionals:
                    for origen in origens_emocionals:
                        doc.add_paragraph(origen)
                else:
                    doc.add_paragraph("No especificat")

                # 📌 Origens nutricionals
                doc.add_heading("Origens nutricionals", level=2)
                origens_nutricionals = origens_adicionals.get(òrgans_definitius[0], {}).get("nutrició", []) if òrgans_definitius else []
                if origens_nutricionals:
                    for origen in origens_nutricionals:
                        doc.add_paragraph(origen)
                else:
                    doc.add_paragraph("No especificat")

                # 📌 Origens de caràcter
                doc.add_heading("Origens de caràcter", level=2)
                origens_caracter = origens_adicionals.get(òrgans_definitius[0], {}).get("caràcter", []) if òrgans_definitius else []
                if origens_caracter:
                    for origen in origens_caracter:
                        doc.add_paragraph(origen)
                else:
                    doc.add_paragraph("No especificat")

                # 📌 Observacions sobre Origens
                doc.add_heading("Observacions sobre Origens", level=3)
                doc.add_paragraph(st.session_state.get("observacions_origens", "No especificat"))
                
                # 📌 Exploració
                doc.add_heading("Exploració", level=2)
                exploracions_resultats = dades_pacient.get("exploracions_resultats", {})
                if exploracions_resultats:
                    for exploracio, resultat in exploracions_resultats.items():
                        doc.add_paragraph(f"{exploracio}: {resultat}")
                else:
                    doc.add_paragraph("No especificat")

                # Observacions d'Exploració
                doc.add_heading("Observacions sobre Exploració", level=3)
                doc.add_paragraph(st.session_state.get("observacions_exploracio", "No especificat"))

                # 📌 12️⃣ Tractaments seleccionats 
                doc.add_heading("Tractaments aplicats", level=2)

                if tractaments_seleccionats:
                    tractaments_unics = set()  # 🔹 Evita duplicats

                    tractaments_filtrats = st.session_state.get("tractaments_seleccionats_usuari", [])

                    for organ, seccio, tractament in tractaments_filtrats:  # Assegura que només recorrem els tractaments seleccionats correctes
                        if (organ, seccio, tractament) not in tractaments_unics:
                            doc.add_paragraph(f"➡ **{organ}** - {tractament}")  # Secció eliminada del text per evitar repeticions visuals
                            tractaments_unics.add((organ, seccio, tractament))  # Afegir a la llista de tractaments ja mostrats

                else:
                    doc.add_paragraph("No s'han seleccionat tractaments.")
                    
                # 📌 Altres tractaments aplicats a la consulta
                doc.add_heading("Altres tractaments aplicats a la consulta", level=2)
                altres_tractaments = dades_pacient.get("altres_tractaments_consulta", "No especificat").split("\n")
                for tractament in altres_tractaments:
                    doc.add_paragraph(tractament.strip())

                # 📌 Altres recomanacions a casa
                doc.add_heading("Altres recomanacions a casa", level=2)
                altres_recomanacions = dades_pacient.get("altres_recomanacions_casa", "No especificat").split("\n")
                for recomanacio in altres_recomanacions:
                    doc.add_paragraph(recomanacio.strip())

                # 💾 Guardar document
                nom_fitxer = f"Historial_{pacient}.docx"
                doc.save(nom_fitxer)
                return nom_fitxer

            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.oxml import OxmlElement
            
            import os

            # Llistar tots els fitxers disponibles
            path_documents = os.path.join(os.getcwd(), "Documents")
            fitxers_disponibles = os.listdir(path_documents) if os.path.exists(path_documents) else []
            print("Fitxers disponibles:", fitxers_disponibles)
            
            import os
            
            # 📌 Carregar documents externs si existeixen
            import os
            from docx import Document

            path_fitoterapia = "/mnt/data/pauta base i instruccions fitoteràpia.docx"
            path_emocions = "/mnt/data/Com un procés emocional o un desajust nutricional poden generar dolors musculoesquelètics.docx"

            doc_fitoterapia = Document(path_fitoterapia) if os.path.exists(path_fitoterapia) else None
            doc_emocions = Document(path_emocions) if os.path.exists(path_emocions) else None
            
            def afegir_contingut_document(doc, fitxer_path):
                """ Afegeix el contingut d'un document Word a un altre mantenint el format. """
                if os.path.exists(fitxer_path):
                    doc_afegit = Document(fitxer_path)

                    for para in doc_afegit.paragraphs:
                        # Afegeix cada paràgraf al nou document mantenint els estils
                        nou_paragraf = doc.add_paragraph()
                        run = nou_paragraf.add_run(para.text)
                        
                        # Manté el format de negreta, cursiva i mida de lletra
                        run.bold = para.runs[0].bold if para.runs else False
                        run.italic = para.runs[0].italic if para.runs else False
                        run.font.size = para.runs[0].font.size if para.runs and para.runs[0].font.size else Pt(10)
                        run.font.name = "Arial"
                    
                    doc.add_paragraph("")  # Afegeix un espai al final del document

                else:
                    print(f"⚠️ No s'ha trobat el fitxer: {fitxer_path}")

            import os
            from docx import Document

            # Definir la ruta dels documents dins la carpeta "documents"
            path_fitoterapia = os.path.join(os.getcwd(), "documents", "pauta base i instruccions fitoteràpia.docx")
            path_emocions = os.path.join(os.getcwd(), "documents", "Com un procés emocional o un desajust nutricional poden generar dolors musculoesquelètics.docx")

            # Comprovar si els fitxers existeixen abans d'intentar carregar-los
            doc_fitoterapia = Document(path_fitoterapia) if os.path.exists(path_fitoterapia) else None
            doc_emocions = Document(path_emocions) if os.path.exists(path_emocions) else None

            def generar_recomanacions_pacient(pacient, tractaments_seleccionats, tractaments):
                """
                Genera un document Word amb només els tractaments seleccionats per al pacient.
                Aplica formatació correcta: Arial, vinyetes, títols diferenciats i enllaços funcionals.
                """

                doc = Document()
                # 🔹 Aplicar Arial 11 com a font base per a cada paràgraf explícitament
                for para in doc.paragraphs:
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)

                # 🔹 Afegir el títol "Recomanacions per a..." amb format correcte
                title = doc.add_paragraph()
                title_run = title.add_run(f"Recomanacions per a {pacient}")
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el títol
                title_run.bold = True
                title_run.font.size = Pt(16)  # Mida gran per al títol
                title_run.font.name = "Arial"
                title_run.font.color.rgb = RGBColor(0, 0, 0)  # Negre

                doc.add_paragraph("")  # Espai després del títol

                # Organitzar tractaments segons la secció
                seccions_ordenades = ["A casa", "Nutrició", "Fitoteràpia", "Fitoteràpia categories"]
                tractaments_organitzats = {seccio: [] for seccio in seccions_ordenades}
                tractaments_unics = set()  # Conjunt per evitar duplicats

                # Organitzar tractaments per secció, incloent Fitoteràpia categories
                # ✅ Assegurar que només es guarden els tractaments realment seleccionats
                for element in st.session_state.get("tractaments_seleccionats_usuari", []):
                    if len(element) == 3:
                        organ, seccio, tractament = element
                    elif len(element) == 4:
                        organ, seccio, categoria, tractament = element  # Captura correctament els 4 valors
                    else:
                        st.write(f"⚠️ Format inesperat per a l'element: {element}")
                        continue  # Continuem amb la següent iteració si el format és incorrecte

                    # ✅ Assegurar que només s'afegeixen tractaments que realment estan seleccionats
                    if (organ, seccio, tractament) in st.session_state["tractaments_seleccionats_usuari"] or \
                       (organ, seccio, categoria, tractament) in st.session_state["tractaments_seleccionats_usuari"]:
                        
                        # Assegurar que el diccionari per cada secció existeix
                        if seccio not in tractaments_organitzats:
                            tractaments_organitzats[seccio] = []

                        tractaments_organitzats[seccio].append((organ, tractament))

                # Afegir cada secció en ordre
                for seccio in seccions_ordenades:
                    if tractaments_organitzats[seccio]:
                        if seccio in ["Fitoteràpia", "Fitoteràpia categories"]:
                            heading_fito = doc.add_paragraph()
                            heading_run_fito = heading_fito.add_run("Fitoteràpia")
                            heading_run_fito.bold = True
                            heading_run_fito.underline = True
                            heading_run_fito.font.size = Pt(13)
                            heading_run_fito.font.name = "Arial"

                            # Ara, per cadascun dels tractaments d'aquesta secció, només afegim el detall
                            for organ, tractament in set(tractaments_organitzats[seccio]):
                                text_tractament = tractaments.get(organ, {}).get(seccio, {}).get(tractament, "Detalls no disponibles")
                                for line in text_tractament.split("\n"):
                                    if line.strip():
                                        p = doc.add_paragraph(style="List Bullet")
                                        run = p.add_run(line.strip())
                                        run.font.name = "Arial"
                                        run.font.size = Pt(11)
                   
                        else:
                            # Per altres seccions, com "A casa" o "Nutrició", mantenim el comportament original
                            if seccio == "Nutrició":
                                heading_dieta = doc.add_paragraph()
                                heading_run_dieta = heading_dieta.add_run("Dieta")
                                heading_run_dieta.bold = True
                                heading_run_dieta.underline = True
                                heading_run_dieta.font.size = Pt(13)
                                heading_run_dieta.font.name = "Arial"

                            # Mostrar cada tractament amb el seu nom com a títol
                            for organ, tractament in set(tractaments_organitzats[seccio]):
                                heading = doc.add_paragraph()
                                heading_run = heading.add_run(tractament)
                                heading_run.bold = True
                                heading_run.font.size = Pt(12)  
                                heading_run.font.name = "Arial"
                                heading_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                text_tractament = tractaments.get(organ, {}).get(seccio, {}).get(tractament, "Detalls no disponibles")
                                for line in text_tractament.split("\n"):
                                    if line.strip():
                                        p = doc.add_paragraph(style="List Bullet")
                                        run = p.add_run(line.strip())
                                        run.font.name = "Arial"
                                        run.font.size = Pt(11)
                                        
                                # 🔹 Afegir la imatge només per "A casa"
                                if seccio == "A casa":
                                    image_filename_png = f"{tractament.replace(' ', '_')}.png"
                                    image_filename_jpg = f"{tractament.replace(' ', '_')}.jpg"

                                    image_path_png = os.path.abspath(os.path.join("imatges", image_filename_png))
                                    image_path_jpg = os.path.abspath(os.path.join("imatges", image_filename_jpg))

                                    if os.path.exists(image_path_png):
                                        doc.add_picture(image_path_png, width=Inches(2.5))
                                    elif os.path.exists(image_path_jpg):
                                        doc.add_picture(image_path_jpg, width=Inches(2.5))

                else:
                    doc.add_paragraph("")

                # 🔹 Afegir secció d'altres recomanacions a casa
                doc.add_paragraph("")
                altres_recomanacions = [rec.strip() for rec in st.session_state.get("altres_recomanacions_casa", "").split("\n") if rec.strip()]

                if altres_recomanacions:  # Només afegir la secció si hi ha contingut
                    heading_recomanacions = doc.add_paragraph()
                    heading_run = heading_recomanacions.add_run("Altres recomanacions a casa")
                    heading_run.bold = True
                    heading_run.font.size = Pt(12)
                    heading_run.font.name = "Arial"

                    for recomanacio in altres_recomanacions:
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(recomanacio)
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                                          
                # 🔹 Funció per afegir el contingut d'un document extern mantenint el format
                def afegir_contingut_document(doc, document_extern):
                    """ Afegeix el contingut d'un document Word a un altre mantenint el format, incloent vinyetes. """
                    if document_extern:
                        for para in document_extern.paragraphs:
                            text = para.text.strip()
                            if not text:
                                continue  # Saltar línies buides
                            
                            # Detectar si el paràgraf té una llista de vinyetes
                            if para.style.name.startswith("List"):
                                nou_paragraf = doc.add_paragraph(style="List Bullet")  # Aplicar estil de llista
                                run = nou_paragraf.add_run(text)
                            else:
                                nou_paragraf = doc.add_paragraph()
                                run = nou_paragraf.add_run(text)

                            # Manté el format original
                            if para.runs:
                                run.bold = para.runs[0].bold
                                run.italic = para.runs[0].italic
                                run.font.size = para.runs[0].font.size if para.runs[0].font.size else Pt(11)
                            
                            # 🔹 Assegurar Arial 11
                            run.font.name = "Arial"
                            run.font.size = Pt(11)

                note = doc.add_paragraph()
                run = note.add_run("⚠️ En cas que algun dels links indicats no funcioni, si us plau, fes-m'ho saber i et buscaré l'enllaç correcte.")
                run.italic = True  # Aplicar cursiva
                run.font.name = "Arial"  # Aplicar Arial
                run.font.size = Pt(11)  # Aplicar mida 11


                # 🔹 Afegir la secció "Instruccions i pauta base de fitoteràpia" si cal
                if doc_fitoterapia and any("Fitoteràpia" in seccio for _, seccio, _ in tractaments_seleccionats):
                    doc.add_page_break()
                    heading = doc.add_paragraph()
                    heading_run = heading.add_run("Instruccions i pauta base de fitoteràpia")
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrat
                    heading_run.bold = True
                    heading_run.font.size = Pt(16)  # Mateixa mida que el títol principal
                    heading_run.font.name = "Arial"
                    heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Negre
                    afegir_contingut_document(doc, doc_fitoterapia)

                # 🔹 Afegir sempre la secció "Com un procés emocional o un desajust nutricional poden generar dolors musculoesquelètics"
                if doc_emocions:
                    doc.add_page_break() 
                    heading = doc.add_paragraph()
                    heading_run = heading.add_run("Com un procés emocional o un desajust nutricional poden generar dolors musculoesquelètics")
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrat
                    heading_run.bold = True
                    heading_run.font.size = Pt(16)  # Mateixa mida que el títol principal
                    heading_run.font.name = "Arial"
                    heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Negre
                    afegir_contingut_document(doc, doc_emocions)

                # 💾 Guardar document
                nom_fitxer = f"Recomanacions_{pacient}.docx"
                doc.save(nom_fitxer)
                return nom_fitxer

            # 📌 Botó per descarregar l'historial de consulta
            if st.button("📄 Generar historial per a la consulta"):
                nom_pacient = st.session_state.get("respostes", {}).get("Nom del pacient", "Sense Nom")
                if nom_pacient != "Sense Nom":
                    nom_fitxer = generar_historial(nom_pacient, st.session_state, st.session_state.get("tractaments_seleccionats_usuari", []))
                    with open(nom_fitxer, "rb") as file:
                        st.download_button("⬇ Descarregar historial", file, file_name=nom_fitxer, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("⚠ Introdueix el nom del pacient abans de generar l'historial.")


            # 📌 Botó per descarregar el tractament per al pacient
            if st.button("📄 Generar recomanacions per al pacient"):
                nom_pacient = st.session_state.get("respostes", {}).get("Nom del pacient", "Sense Nom")
                if nom_pacient != "Sense Nom":
                    nom_fitxer = generar_recomanacions_pacient(nom_pacient, st.session_state.get("tractaments_seleccionats_usuari", []), tractaments)
                    with open(nom_fitxer, "rb") as file:
                        st.download_button("⬇ Descarregar tractament", file, file_name=nom_fitxer, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("⚠ Introdueix el nom del pacient abans de generar el document.")

if __name__ == "__main__":
    assistent = AssistentDolencies()
    assistent.iniciar_qüestionari()